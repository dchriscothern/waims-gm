from __future__ import annotations

import csv
import json
import os
import sys
from datetime import datetime, timezone
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional
from uuid import uuid4

ROOT = Path(__file__).resolve().parent
VENV_DIR = ROOT / ".venv"
VENV_PYTHON = VENV_DIR / "Scripts" / "python.exe"


def _using_repo_venv() -> bool:
    current_exec = Path(sys.executable).resolve()
    return VENV_DIR.resolve() in current_exec.parents or current_exec == VENV_PYTHON.resolve()


def _relaunch_with_repo_venv() -> None:
    if not VENV_PYTHON.exists():
        return
    if _using_repo_venv():
        return
    if os.environ.get("WAIMS_GM_BOOTSTRAPPED") == "1":
        return

    os.environ["WAIMS_GM_BOOTSTRAPPED"] = "1"
    orig_argv = list(getattr(sys, "orig_argv", []) or [])
    if len(orig_argv) >= 2:
        relaunch_argv = [str(VENV_PYTHON), *orig_argv[1:]]
    else:
        relaunch_argv = [str(VENV_PYTHON), "-m", "streamlit", "run", str(ROOT / "streamlit_app.py")]
    os.execv(str(VENV_PYTHON), relaunch_argv)


_relaunch_with_repo_venv()

import httpx
import streamlit as st
from app.config import API_BASE_URL, IS_LIVE_ENV, WAIMS_DEMO_MODE, WAIMS_ENV_LABEL, WAIMS_PYTHON_BASE_URL
from waims_gm.domain import Player, TeamContext
from waims_gm.demo_data import demo_payloads
from waims_gm.services import evaluate_single_player
from waims_gm_recruiting import recruiting_tab

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    WORD_EXPORT_AVAILABLE = True
    WORD_EXPORT_ERROR = ""
except Exception as e:
    WORD_EXPORT_AVAILABLE = False
    WORD_EXPORT_ERROR = str(e)

st.set_page_config(
    page_title="WAIMS-GM",
    page_icon=":basketball:",
    layout="wide",
    initial_sidebar_state="expanded",
)

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=IBM+Plex+Mono:wght@400;500&family=Libre+Baskerville:wght@400;700&display=swap');

:root {
    --bg: #f5f0e6;
    --paper: #f8f3ea;
    --ink: #161616;
    --muted: #5f5a52;
    --gold: #9f7a2d;
    --line: rgba(22,22,22,0.12);
    --sidebar: #131313;
    --sidebar-ink: #e8dfd1;
    --card: rgba(255,255,255,0.52);
    --success: #345c3e;
    --danger: #8a3d32;
}

html, body, [data-testid="stAppViewContainer"] {
    background:
        linear-gradient(to bottom, rgba(0,0,0,0.015), rgba(0,0,0,0.015)),
        repeating-linear-gradient(
            0deg,
            rgba(0,0,0,0.00) 0px,
            rgba(0,0,0,0.00) 22px,
            rgba(0,0,0,0.018) 23px
        ),
        var(--bg);
    color: var(--ink);
    font-family: "Libre Baskerville", serif;
}

[data-testid="stSidebar"] {
    background: var(--sidebar);
    color: var(--sidebar-ink);
    border-right: 1px solid rgba(255,255,255,0.06);
}

[data-testid="stSidebar"] * {
    color: var(--sidebar-ink);
}

.block-container {
    padding-top: 1.2rem;
    padding-bottom: 2rem;
    max-width: 1720px;
}

.waims-header {
    border-top: 3px solid var(--gold);
    border-bottom: 1px solid var(--line);
    padding: 0.6rem 0 1rem 0;
    margin-bottom: 1rem;
}

.waims-kicker {
    font-family: "IBM Plex Mono", monospace;
    text-transform: uppercase;
    letter-spacing: 0.12em;
    font-size: 0.72rem;
    color: var(--muted);
    margin-bottom: 0.35rem;
}

.waims-title {
    font-family: "Playfair Display", serif;
    font-size: 2.2rem;
    line-height: 1.05;
    color: var(--ink);
    margin: 0;
}

.waims-subtitle {
    margin-top: 0.45rem;
    color: var(--muted);
    font-size: 0.98rem;
}

.waims-meta-row {
    margin-top: 0.7rem;
    display: flex;
    gap: 0.6rem;
    align-items: center;
    flex-wrap: wrap;
}

.env-badge {
    display: inline-flex;
    align-items: center;
    border-radius: 999px;
    border: 1px solid var(--line);
    padding: 0.18rem 0.65rem;
    font-family: "IBM Plex Mono", monospace;
    font-size: 0.72rem;
    letter-spacing: 0.04em;
    text-transform: uppercase;
}

.env-badge-live {
    background: rgba(138, 61, 50, 0.10);
    color: var(--danger);
    border-color: rgba(138, 61, 50, 0.28);
}

.env-badge-sandbox {
    background: rgba(52, 92, 62, 0.10);
    color: var(--success);
    border-color: rgba(52, 92, 62, 0.28);
}

.metric-card {
    background: var(--card);
    border: 1px solid var(--line);
    border-top: 2px solid var(--gold);
    padding: 0.9rem 1rem;
    border-radius: 10px;
    min-height: 118px;
}

.metric-label {
    font-family: "IBM Plex Mono", monospace;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    font-size: 0.72rem;
    color: var(--muted);
}

.metric-value {
    font-family: "Playfair Display", serif;
    font-size: 1.85rem;
    margin-top: 0.2rem;
    color: var(--ink);
}

.metric-note {
    margin-top: 0.3rem;
    font-size: 0.85rem;
    color: var(--muted);
}

.section-title {
    font-family: "Playfair Display", serif;
    font-size: 1.28rem;
    margin: 1rem 0 0.5rem 0;
    padding-bottom: 0.3rem;
    border-bottom: 1px solid var(--line);
}

.section-kicker {
    font-family: "IBM Plex Mono", monospace;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    font-size: 0.72rem;
    color: var(--muted);
    margin-bottom: 0.25rem;
}

.intake-shell,
.detail-shell,
.board-card,
.soft-card {
    background: rgba(255,255,255,0.55);
    border: 1px solid var(--line);
    border-radius: 12px;
    width: 100%;
    min-width: 0;
    box-sizing: border-box;
}

.intake-shell,
.detail-shell {
    border-top: 2px solid var(--gold);
    padding: 1rem;
}

.intake-blurb {
    color: var(--muted);
    font-size: 0.93rem;
    margin-bottom: 0.8rem;
}

.board-card {
    border-left: 3px solid var(--gold);
    padding: 0.85rem 1rem;
    margin-bottom: 0.75rem;
}

.board-card-selected {
    background: rgba(255,255,255,0.76);
    border: 1px solid rgba(159,122,45,0.45);
    border-left: 4px solid var(--gold);
    box-shadow: 0 2px 10px rgba(0,0,0,0.04);
}

.board-head {
    display: flex;
    justify-content: space-between;
    gap: 1rem;
    align-items: baseline;
    flex-wrap: wrap;
}

.board-name {
    font-family: "Playfair Display", serif;
    font-size: 1.15rem;
    color: var(--ink);
    min-width: 0;
    overflow-wrap: anywhere;
}

.board-tag {
    font-family: "IBM Plex Mono", monospace;
    font-size: 0.72rem;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: var(--gold);
}

.board-meta {
    margin-top: 0.35rem;
    color: var(--muted);
    font-size: 0.9rem;
}

.board-note {
    margin-top: 0.45rem;
    color: var(--ink);
    font-size: 0.88rem;
    line-height: 1.45;
}

.player-title {
    font-family: "Playfair Display", serif;
    font-size: clamp(1.4rem, 2.8vw, 1.7rem);
    color: var(--ink);
    margin-bottom: 0.15rem;
    overflow-wrap: anywhere;
}

.player-meta {
    color: var(--muted);
    font-size: 0.95rem;
    overflow-wrap: anywhere;
}

.diagnostic-mini-card {
    border: 1px solid rgba(159,122,45,0.18);
    border-radius: 14px;
    padding: 0.5rem 0.65rem;
    background: rgba(255,255,255,0.72);
    min-width: 0;
}

.diagnostic-mini-label {
    font-family: "IBM Plex Mono", monospace;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    font-size: 0.62rem;
    color: var(--muted);
    margin-bottom: 0.18rem;
}

.diagnostic-mini-grade {
    font-family: "Playfair Display", serif;
    font-size: 1.1rem;
    line-height: 1;
    color: var(--ink);
}

.diagnostic-row-head {
    display: flex;
    align-items: center;
    gap: 0.7rem;
    margin-bottom: 0.35rem;
}

.diagnostic-row-grade {
    font-family: "Playfair Display", serif;
    font-size: 1.35rem;
    line-height: 1;
    color: var(--gold);
    min-width: 1.2rem;
}

.rule {
    border-top: 1px solid var(--line);
    margin: 0.95rem 0 1rem 0;
}

.soft-card {
    padding: 0.9rem 1rem;
    height: 100%;
    overflow: hidden;
}

.mini-label {
    font-family: "IBM Plex Mono", monospace;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    font-size: 0.7rem;
    color: var(--muted);
    margin-bottom: 0.35rem;
}

.memo-text {
    color: var(--ink);
    line-height: 1.6;
    font-size: 0.97rem;
    overflow-wrap: anywhere;
    word-break: break-word;
}

.subtle-list {
    margin: 0.25rem 0 0 0;
    padding-left: 1.1rem;
}

.subtle-list li {
    margin-bottom: 0.32rem;
    color: var(--ink);
}

.filter-note {
    color: var(--muted);
    font-size: 0.84rem;
}

.action-draft { color: var(--success); }
.action-sign { color: var(--gold); }
.action-pass { color: var(--danger); }

.score-grid {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(165px, 1fr));
    gap: 0.75rem;
    margin-bottom: 0.85rem;
    align-items: stretch;
}

.profile-grid {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(145px, 1fr));
    gap: 0.75rem;
}

.profile-card {
    background: rgba(255,255,255,0.50);
    border: 1px solid var(--line);
    border-radius: 12px;
    padding: 0.8rem 0.9rem;
    min-height: 108px;
    width: 100%;
    min-width: 0;
    box-sizing: border-box;
}

.profile-label {
    color: var(--muted);
    font-size: 0.95rem;
    margin-bottom: 0.35rem;
    overflow-wrap: anywhere;
    word-break: break-word;
}

.profile-value {
    font-family: "Playfair Display", serif;
    font-size: clamp(1.9rem, 3.4vw, 3rem);
    line-height: 1.02;
    color: var(--ink);
    overflow-wrap: anywhere;
    word-break: break-word;
}

.score-card {
    background: rgba(255,255,255,0.50);
    border: 1px solid var(--line);
    border-radius: 12px;
    padding: 0.8rem 0.9rem;
    min-height: 108px;
    width: 100%;
    min-width: 0;
    box-sizing: border-box;
}

.score-card-wide {
    grid-column: 1 / -1;
}

.score-label {
    font-size: 0.95rem;
    color: var(--muted);
    margin-bottom: 0.35rem;
}

.score-value {
    font-family: "Playfair Display", serif;
    font-size: clamp(1.65rem, 3.1vw, 2.35rem);
    line-height: 1.05;
    color: var(--ink);
    white-space: normal;
    overflow-wrap: anywhere;
    word-break: break-word;
}

textarea {
    min-height: 110px !important;
}

.stButton > button {
    border-radius: 8px;
}

div[data-testid="stMetric"] {
    border: 1px solid var(--line);
    border-radius: 10px;
    background: rgba(255,255,255,0.42);
    padding: 0.45rem 0.7rem;
    min-width: 0;
    width: 100%;
    box-sizing: border-box;
    overflow: hidden;
}

div[data-testid="stMetricLabel"] {
    overflow-wrap: anywhere;
    word-break: break-word;
    min-width: 0;
}

div[data-testid="stMetricValue"] {
    font-size: clamp(1.6rem, 3vw, 2.35rem);
    line-height: 1.05;
    overflow-wrap: anywhere;
    word-break: break-word;
    white-space: normal;
    min-width: 0;
}

div[data-testid="stMetricLabel"] > div,
div[data-testid="stMetricValue"] > div {
    overflow-wrap: anywhere;
    word-break: break-word;
    white-space: normal;
}

[data-testid="column"] {
    min-width: 0;
}

hr {
    border: none;
    border-top: 1px solid var(--line);
    margin: 0.8rem 0 1rem 0;
}

.compare-grid {
    display: grid;
    gap: 0.55rem;
    min-width: 0;
}

.compare-row {
    display: grid;
    grid-template-columns: minmax(110px, 1.2fr) minmax(80px, 0.9fr) minmax(90px, 0.8fr) minmax(80px, 0.9fr);
    gap: 0.6rem;
    align-items: center;
    background: rgba(255,255,255,0.45);
    border: 1px solid var(--line);
    border-radius: 10px;
    padding: 0.7rem 0.85rem;
    min-width: 0;
}

.compare-metric {
    font-family: "IBM Plex Mono", monospace;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    font-size: 0.72rem;
    color: var(--muted);
}

.compare-score {
    font-family: "Playfair Display", serif;
    font-size: 1.18rem;
    color: var(--ink);
    overflow-wrap: anywhere;
}

.compare-advantage {
    text-align: center;
    color: var(--muted);
    font-size: 0.82rem;
    overflow-wrap: anywhere;
}

@media (max-width: 1200px) {
    .compare-row {
        grid-template-columns: repeat(2, minmax(0, 1fr));
    }
}
</style>
"""

MODE_LABELS = {
    "pro_wnba": "Pro / WNBA",
    "cbb_high_major": "CBB High-Major",
    "cbb_d2_low_resource": "CBB D2-3 NAIA Juco",
    "recruiting_only": "Recruiting Only",
}

WORKFLOW_STAGES = ["Shadow Board", "Priority", "Live Board", "Closed"]
MEDICAL_DILIGENCE_LABEL = "Med Diligence"
MEDICAL_DILIGENCE_LEVELS = [
    "Clear for pursuit",
    "Proceed with caution",
    "Needs deeper diligence",
    "High risk file",
]
MEDICAL_CONFIDENCE_LEVELS = ["Low", "Medium", "High"]
WORKSPACE_ROLES = {
    "gm": "General Manager",
    "sport_science": "Sport Science",
    "medical": "Medical",
}
COLLABORATOR_ROLES = {"sport_science", "medical"}
WORKSPACE_LOGIN_USERS = {
    "gm": {"password": "gm123", "role": "gm", "label": "General Manager"},
    "scientist": {"password": "sci123", "role": "sport_science", "label": "Sport Science"},
    "medical": {"password": "med123", "role": "medical", "label": "Medical"},
}

DEFAULT_MODE = "cbb_d2_low_resource"

MODE_PLAYBOOKS: Dict[str, Dict[str, str]] = {
    "pro_wnba": {
        "headline": "Pro / WNBA acquisition lens",
        "primary_question": "Can this player solve a rotation problem without creating a value leak against pro-level cost and durability pressure?",
        "weight_note": "Impact drives the file first, then availability and role fit.",
        "create_note": "Use this mode when you care about immediate pro translation, contract pressure, and real rotation minutes.",
        "board_note": "Board view favors playable impact and cleaner short-window dependability over speculative upside.",
        "dossier_note": "Read the dossier as a front-office acquisition file with immediate rotation and cost discipline in mind.",
        "compare_note": "Compare mode should answer which player solves the current rotation question with the cleaner value profile.",
    },
    "cbb_high_major": {
        "headline": "High-major roster-building lens",
        "primary_question": "Can this player help a high-major program win quickly while still justifying portal cost and role usage?",
        "weight_note": "Impact and upside stay aggressive, with fit and value still important but not dominant.",
        "create_note": "Use this mode for portal leads, top-end adds, and players who must hold up against high-major role pressure.",
        "board_note": "Board view leans toward immediate game impact, usage scalability, and portal translation.",
        "dossier_note": "Read the dossier as a high-major add file: immediate usefulness matters, but a ceiling signal still moves the recommendation.",
        "compare_note": "Compare mode should tell you who is the better portal bet for your current high-major rotation and offense ecosystem.",
    },
    "cbb_d2_low_resource": {
        "headline": "D2 / NAIA / JUCO roster-build lens",
        "primary_question": "Can this player help us win soon without forcing our staff into a cost, risk, or roster-fit compromise?",
        "weight_note": "Value, fit, and availability matter more here than pure upside chasing.",
        "create_note": "Use this mode for D2, D3, NAIA, and Juco-style decisions where affordability, role clarity, and readiness matter.",
        "board_note": "Board view favors ready, affordable players who solve a live roster need without introducing avoidable instability.",
        "dossier_note": "Read the dossier as a small-staff roster file: the best target is often the one who helps now with the least operational drag.",
        "compare_note": "Compare mode should answer who gives the cleaner immediate return for the staff and roster you actually have.",
    },
    "recruiting_only": {
        "headline": "Long-horizon recruiting lens",
        "primary_question": "Is this prospect worth betting on because the long-term upside and age-adjusted runway justify the uncertainty?",
        "weight_note": "Upside and future value matter more than immediate impact or current minutes stability.",
        "create_note": "Use this mode for high-school, prep, or long-range recruiting decisions where runway matters more than instant rotation help.",
        "board_note": "Board view favors growth curve, positional need, and long-term upside more than near-term certainty.",
        "dossier_note": "Read the dossier as a projection file. The recommendation should reflect future leverage, not just current readiness.",
        "compare_note": "Compare mode should tell you which prospect is the better long-horizon bet for your program's future roster shape.",
    },
}

ACTION_LABELS = {
    "pro_wnba": {"draft": "Draft", "sign": "Sign", "pass": "Pass"},
    "cbb_high_major": {"draft": "Priority Target", "sign": "Take/Add", "pass": "Pass"},
    "cbb_d2_low_resource": {"draft": "Priority Target", "sign": "Fit Add", "pass": "Pass"},
    "recruiting_only": {"draft": "High Priority", "sign": "Monitor", "pass": "Pass"},
}

COMPONENT_LABELS = {
    "fit": "Fit",
    "impact": "Impact",
    "upside": "Upside",
    "availability": "Availability",
    "value": "Value",
}

CSV_IMPORT_COLUMNS = [
    "display_name",
    "player_id",
    "player_name",
    "position",
    "age",
    "offense_rating",
    "defense_rating",
    "shooting_rating",
    "playmaking_rating",
    "rebounding_rating",
    "health_risk",
    "upside",
    "minutes_stability",
    "expected_cost_tier",
    "team_id",
    "timeline",
    "need_g",
    "need_f",
    "need_c",
    "cap_flexibility",
    "risk_tolerance",
    "summary_note",
    "strengths",
    "concerns",
    "mode",
]

CSV_REQUIRED_COLUMNS = [
    "player_id",
    "player_name",
    "position",
    "age",
    "offense_rating",
    "defense_rating",
    "shooting_rating",
    "playmaking_rating",
    "rebounding_rating",
    "health_risk",
    "upside",
    "minutes_stability",
    "expected_cost_tier",
    "team_id",
    "timeline",
    "need_g",
    "need_f",
    "need_c",
    "cap_flexibility",
    "risk_tolerance",
]

EXAMPLE_IMPORT_CSV_PATH = Path(__file__).resolve().parent / "examples" / "waims_gm_import_sample.csv"
WAIMS_PYTHON_HANDOFF_DIR = Path(__file__).resolve().parent / "runtime" / "waims_python_handoffs"
PROSPECT_REQUEST_DIR = Path(__file__).resolve().parent / "runtime" / "prospect_diligence_requests"

PRESETS: Dict[str, Dict[str, Dict[str, Any]]] = {
    "pro_wnba": {
        "3-and-D Wing": {
            "position": "F",
            "age": 24,
            "offense_rating": 72.0,
            "defense_rating": 82.0,
            "shooting_rating": 78.0,
            "playmaking_rating": 58.0,
            "rebounding_rating": 68.0,
            "health_risk": 0.20,
            "upside": 0.70,
            "minutes_stability": 0.78,
            "expected_cost_tier": 2,
            "need_g": 0.40,
            "need_f": 0.85,
            "need_c": 0.35,
            "summary_note": "Reliable wing target who fits a low-maintenance rotation role.",
            "strengths": "Defends multiple positions\nReliable catch-and-shoot spacing\nStable role acceptance",
            "concerns": "Limited self-creation\nMay be capped as a secondary piece",
        },
        "Second Unit Pace Guard": {
            "position": "G",
            "age": 26,
            "offense_rating": 75.0,
            "defense_rating": 61.0,
            "shooting_rating": 74.0,
            "playmaking_rating": 77.0,
            "rebounding_rating": 34.0,
            "health_risk": 0.18,
            "upside": 0.58,
            "minutes_stability": 0.74,
            "expected_cost_tier": 2,
            "need_g": 0.72,
            "need_f": 0.24,
            "need_c": 0.16,
            "summary_note": "Bench guard target who can stabilize second-unit offense without premium acquisition cost.",
            "strengths": "Pick-and-roll command\nSecond-unit shot creation\nTempo control",
            "concerns": "Size pressure on defense\nLess upside than younger guard bets",
        },
    },
    "cbb_high_major": {
        "Portal Lead Guard": {
            "position": "G",
            "age": 21,
            "offense_rating": 80.0,
            "defense_rating": 68.0,
            "shooting_rating": 76.0,
            "playmaking_rating": 84.0,
            "rebounding_rating": 46.0,
            "health_risk": 0.16,
            "upside": 0.76,
            "minutes_stability": 0.79,
            "expected_cost_tier": 3,
            "need_g": 0.85,
            "need_f": 0.45,
            "need_c": 0.20,
            "summary_note": "High-major portal guard with immediate usage and ball-screen value.",
            "strengths": "Creation\nDecision-making\nPull-up threat\nLate-clock offense",
            "concerns": "Defensive translation\nPrice sensitivity\nBall-dominant profile",
        },
        "Portal Two-Way Wing": {
            "position": "F",
            "age": 22,
            "offense_rating": 76.0,
            "defense_rating": 79.0,
            "shooting_rating": 77.0,
            "playmaking_rating": 61.0,
            "rebounding_rating": 63.0,
            "health_risk": 0.13,
            "upside": 0.73,
            "minutes_stability": 0.80,
            "expected_cost_tier": 3,
            "need_g": 0.35,
            "need_f": 0.84,
            "need_c": 0.22,
            "summary_note": "Portal wing with high-major defensive portability and enough shooting to hold a real role.",
            "strengths": "Positional defense\nCatch-and-shoot value\nRole flexibility",
            "concerns": "May not carry primary creation\nPortal market could drive price",
        },
    },
    "cbb_d2_low_resource": {
        "D2 Two-Way Wing": {
            "position": "F",
            "age": 20,
            "offense_rating": 71.0,
            "defense_rating": 77.0,
            "shooting_rating": 74.0,
            "playmaking_rating": 55.0,
            "rebounding_rating": 67.0,
            "health_risk": 0.15,
            "upside": 0.72,
            "minutes_stability": 0.81,
            "expected_cost_tier": 1,
            "need_g": 0.35,
            "need_f": 0.82,
            "need_c": 0.34,
            "summary_note": "Affordable two-way wing who can help a smaller staff win quickly.",
            "strengths": "Role clarity\nDefensive versatility\nShooting floor\nStrong budget fit",
            "concerns": "Creation ceiling\nMay need usage protection",
        },
        "Ready Guard Value Add": {
            "position": "G",
            "age": 22,
            "offense_rating": 73.0,
            "defense_rating": 74.0,
            "shooting_rating": 70.0,
            "playmaking_rating": 72.0,
            "rebounding_rating": 38.0,
            "health_risk": 0.11,
            "upside": 0.66,
            "minutes_stability": 0.86,
            "expected_cost_tier": 1,
            "need_g": 0.86,
            "need_f": 0.34,
            "need_c": 0.20,
            "summary_note": "Ready-made backcourt add who gives a smaller staff immediate competence and cheap minutes.",
            "strengths": "Decision security\nRole readiness\nBudget-friendly fit",
            "concerns": "Lower long-term ceiling\nDoes not solve frontcourt size",
        },
    },
    "recruiting_only": {
        "High-Upside Combo Guard": {
            "position": "G",
            "age": 18,
            "offense_rating": 73.0,
            "defense_rating": 63.0,
            "shooting_rating": 75.0,
            "playmaking_rating": 71.0,
            "rebounding_rating": 39.0,
            "health_risk": 0.12,
            "upside": 0.84,
            "minutes_stability": 0.60,
            "expected_cost_tier": 2,
            "need_g": 0.75,
            "need_f": 0.30,
            "need_c": 0.15,
            "summary_note": "Developmental guard with real long-term upside.",
            "strengths": "Shot-making\nHandle creativity\nGrowth curve",
            "concerns": "Physical maturity\nDecision stability\nDefensive readiness",
        },
        "Long Wing Upside Bet": {
            "position": "F",
            "age": 17,
            "offense_rating": 67.0,
            "defense_rating": 64.0,
            "shooting_rating": 71.0,
            "playmaking_rating": 56.0,
            "rebounding_rating": 61.0,
            "health_risk": 0.09,
            "upside": 0.89,
            "minutes_stability": 0.52,
            "expected_cost_tier": 2,
            "need_g": 0.28,
            "need_f": 0.81,
            "need_c": 0.26,
            "summary_note": "Length-and-skill wing bet with long-term lineup versatility if the physical development comes.",
            "strengths": "Frame upside\nProjectable jumper\nTwo-way runway",
            "concerns": "Needs strength\nCurrent impact is limited",
        },
    },
}

TEAM_CONTEXT_PRESETS: Dict[str, Dict[str, Dict[str, Any]]] = {
    "pro_wnba": {
        "Rotation Wing Need": {
            "team_id": "pro-team-1",
            "timeline": "win_now",
            "cap_flexibility": 0.42,
            "risk_tolerance": 0.34,
            "need_g": 0.38,
            "need_f": 0.84,
            "need_c": 0.28,
        },
    },
    "cbb_high_major": {
        "Portal Lead Guard Need": {
            "team_id": "hm-team-1",
            "timeline": "win_now",
            "cap_flexibility": 0.56,
            "risk_tolerance": 0.42,
            "need_g": 0.86,
            "need_f": 0.44,
            "need_c": 0.22,
        },
    },
    "cbb_d2_low_resource": {
        "Small-Staff Wing Priority": {
            "team_id": "d2-team-1",
            "timeline": "balanced",
            "cap_flexibility": 0.64,
            "risk_tolerance": 0.38,
            "need_g": 0.46,
            "need_f": 0.82,
            "need_c": 0.32,
        },
        "Immediate Guard Value Add": {
            "team_id": "d2-team-1",
            "timeline": "balanced",
            "cap_flexibility": 0.68,
            "risk_tolerance": 0.34,
            "need_g": 0.88,
            "need_f": 0.34,
            "need_c": 0.18,
        },
    },
    "recruiting_only": {
        "Future Wing Build": {
            "team_id": "recruit-team-1",
            "timeline": "rebuild",
            "cap_flexibility": 0.58,
            "risk_tolerance": 0.52,
            "need_g": 0.32,
            "need_f": 0.80,
            "need_c": 0.24,
        },
    },
}


def inject_css() -> None:
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


def api_headers(token: str) -> Dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _market_band_from_cost_tier(cost_tier: Any) -> str:
    cost = _safe_float(cost_tier)
    if cost <= 1:
        return "Budget"
    if cost <= 3:
        return "Value"
    if cost <= 6:
        return "Market"
    return "Premium"


def _priority_tier_from_score(score: Any) -> str:
    numeric = _safe_float(score)
    if numeric >= 80:
        return "Tier 1"
    if numeric >= 72:
        return "Tier 2"
    return "Tier 3"


def _team_preset_store() -> Dict[str, Dict[str, Dict[str, Any]]]:
    if "team_context_presets" not in st.session_state:
        st.session_state["team_context_presets"] = {}
    return st.session_state["team_context_presets"]


def get_team_context_presets(mode: str) -> Dict[str, Dict[str, Any]]:
    presets = dict(TEAM_CONTEXT_PRESETS.get(mode, {}))
    presets.update(_team_preset_store().get(mode, {}))
    return presets


def save_team_context_preset(mode: str, name: str, values: Dict[str, Any]) -> None:
    cleaned_name = name.strip()
    if not cleaned_name:
        return
    store = _team_preset_store()
    mode_store = dict(store.get(mode) or {})
    mode_store[cleaned_name] = {
        "team_id": str(values.get("team_id") or ""),
        "timeline": str(values.get("timeline") or "balanced"),
        "cap_flexibility": float(values.get("cap_flexibility") or 0),
        "risk_tolerance": float(values.get("risk_tolerance") or 0),
        "need_g": float(values.get("need_g") or 0),
        "need_f": float(values.get("need_f") or 0),
        "need_c": float(values.get("need_c") or 0),
    }
    store[mode] = mode_store


def build_med_diligence_handoff(detail: Dict[str, Any]) -> Dict[str, Any]:
    detail = normalize_detail_for_display(detail) if "components" in detail else dict(detail)
    workflow = get_front_office_meta(detail)
    player = detail.get("player", {}) or {}
    return {
        "schema_version": "waims_med_handoff_v1",
        "source_app": "waims_gm",
        "evaluation_id": str(detail.get("id") or ""),
        "player_id": str(player.get("id") or ""),
        "player_name": str(player.get("name") or ""),
        "team_id": str(detail.get("team_id") or (detail.get("ctx", {}) or {}).get("team_id") or ""),
        "med_diligence": {
            "call": workflow.get("level", "Proceed with caution"),
            "confidence": workflow.get("confidence", "Medium"),
            "movement_flag": workflow.get("movement_flag", ""),
            "public_history": workflow.get("public_history", ""),
            "gm_note": workflow.get("gm_note", ""),
            "reviewed_by": workflow.get("reviewed_by", workflow.get("owner", "GM")),
            "reviewed_at": workflow.get("reviewed_at", _now_iso()),
            "source": workflow.get("source", "WAIMS-GM manual overlay"),
        },
    }


def build_med_diligence_template() -> str:
    return json.dumps(
        {
            "schema_version": "waims_med_handoff_v1",
            "source_app": "waims_python",
            "evaluation_id": "",
            "player_id": "",
            "player_name": "",
            "team_id": "",
            "med_diligence": {
                "call": "Proceed with caution",
                "confidence": "Medium",
                "movement_flag": "",
                "public_history": "",
                "gm_note": "",
                "reviewed_by": "",
                "reviewed_at": "",
                "source": "WAIMS Python handoff",
            },
        },
        indent=2,
    )


def parse_med_diligence_handoff(raw_text: str) -> Dict[str, str]:
    try:
        payload = json.loads(raw_text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid Med Diligence handoff JSON: {exc}") from exc

    med_block = payload.get("med_diligence") or {}
    level = str(med_block.get("call") or "").strip()
    if level not in MEDICAL_DILIGENCE_LEVELS:
        raise ValueError(f"`med_diligence.call` must be one of {', '.join(MEDICAL_DILIGENCE_LEVELS)}.")

    confidence = str(med_block.get("confidence") or "Medium").strip()
    if confidence not in MEDICAL_CONFIDENCE_LEVELS:
        raise ValueError(f"`med_diligence.confidence` must be one of {', '.join(MEDICAL_CONFIDENCE_LEVELS)}.")

    return {
        "level": level,
        "confidence": confidence,
        "movement_flag": str(med_block.get("movement_flag") or "").strip(),
        "public_history": str(med_block.get("public_history") or "").strip(),
        "gm_note": str(med_block.get("gm_note") or "").strip(),
        "reviewed_by": str(med_block.get("reviewed_by") or "").strip(),
        "reviewed_at": str(med_block.get("reviewed_at") or "").strip() or _now_iso(),
        "source": str(med_block.get("source") or payload.get("source_app") or "WAIMS Python handoff").strip(),
    }


def sync_waims_python_handoffs(evaluations: List[Dict[str, Any]]) -> Dict[str, int]:
    summary = {"files": 0, "matched": 0, "ignored": 0}
    if not WAIMS_PYTHON_HANDOFF_DIR.exists():
        return summary

    by_eval_id = {str(row.get("id") or ""): str(row.get("id") or "") for row in evaluations if row.get("id")}
    by_player_team = {
        (
            str((row.get("player") or {}).get("id") or "").strip().lower(),
            str(row.get("team_id") or "").strip().lower(),
        ): str(row.get("id") or "")
        for row in evaluations
        if (row.get("player") or {}).get("id") and row.get("team_id")
    }

    for path in sorted(WAIMS_PYTHON_HANDOFF_DIR.glob("*.json")):
        summary["files"] += 1
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
            imported = parse_med_diligence_handoff(json.dumps(payload))
        except Exception:
            summary["ignored"] += 1
            continue

        evaluation_id = str(payload.get("evaluation_id") or "").strip()
        matched_id = by_eval_id.get(evaluation_id)
        if not matched_id:
            key = (
                str(payload.get("player_id") or "").strip().lower(),
                str(payload.get("team_id") or "").strip().lower(),
            )
            matched_id = by_player_team.get(key)

        if not matched_id:
            summary["ignored"] += 1
            continue

        save_front_office_meta(matched_id, **imported)
        summary["matched"] += 1

    return summary


def build_prospect_diligence_request(detail: Dict[str, Any], overrides: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    detail = normalize_detail_for_display(detail) if "components" in detail else dict(detail)
    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    workflow = get_front_office_meta(detail)
    payload = {
        "schema_version": "waims_prospect_request_v1",
        "source_app": "waims_gm",
        "evaluation_id": str(detail.get("id") or ""),
        "player_id": str(player.get("id") or ""),
        "player_name": str(player.get("name") or ""),
        "position": str(player.get("position") or ""),
        "team_id": str(detail.get("team_id") or ctx.get("team_id") or ""),
        "mode": str(detail.get("mode") or DEFAULT_MODE),
        "requested_by": str(workflow.get("owner") or "GM"),
        "priority_tier": str(workflow.get("priority_tier") or "Tier 3"),
        "request_focus": "External prospect diligence",
        "questions": (
            "What public-file readiness, load, injury, and durability concerns should materially change pricing or pursuit?"
        ),
        "gm_context": str(detail.get("summary_note") or summarize_context(ctx)),
    }
    if overrides:
        payload.update({key: value for key, value in overrides.items() if value is not None})
    return payload


def derive_front_office_profile(detail: Dict[str, Any]) -> Dict[str, str]:
    detail = normalize_detail_for_display(detail) if "components" in detail else dict(detail)
    player = detail.get("player", {}) or {}
    mode = detail.get("mode") or DEFAULT_MODE
    action = str(detail.get("recommended_action") or "").lower()
    score = _safe_float(detail.get("overall_score"))
    value_component = _safe_float((detail.get("components") or {}).get("value"), score)
    cost_tier = player.get("expected_cost_tier", 0)
    market_band = _market_band_from_cost_tier(cost_tier)
    priority_tier = _priority_tier_from_score(score)

    if action == "draft" and score >= 80:
        stage = "Live Board"
    elif action in {"draft", "sign"} and score >= 72:
        stage = "Priority"
    else:
        stage = "Shadow Board"

    owner_defaults = {
        "pro_wnba": "GM",
        "cbb_high_major": "GM",
        "cbb_d2_low_resource": "Head Coach / GM",
        "recruiting_only": "Recruiting Lead",
    }
    owner = owner_defaults.get(mode, "GM")

    if value_component >= 76 and _safe_float(cost_tier) <= 2:
        value_posture = "Budget win"
    elif value_component >= 70:
        value_posture = "Fair price"
    elif _safe_float(cost_tier) >= 5:
        value_posture = "Overpay risk"
    else:
        value_posture = "Price-sensitive"

    stage_actions = {
        "Shadow Board": "Keep the file live and gather portal intel before the window opens.",
        "Priority": "Line up staff conviction, contact plan, and a clean value walk-away number.",
        "Live Board": "Move from evaluation to close: confirm fit, price, and decision owner now.",
        "Closed": "Archive the file and note the reason so the board stays clean.",
    }

    spend_note = {
        "Budget": "Works best as a low-cost add or scholarship-efficient target.",
        "Value": "Worth pursuing if the market stays disciplined.",
        "Market": "Playable target, but only if the role is worth the spend.",
        "Premium": "Needs top-end conviction or scarce-positional value to justify the cost.",
    }[market_band]

    return {
        "stage": stage,
        "owner": owner,
        "next_action": stage_actions[stage],
        "market_band": market_band,
        "value_posture": value_posture,
        "spend_discipline": spend_note,
        "priority_tier": priority_tier,
    }


def derive_medical_diligence_profile(detail: Dict[str, Any]) -> Dict[str, str]:
    return {
        "level": "Proceed with caution",
        "confidence": "Low",
        "movement_flag": "",
        "public_history": "",
        "gm_note": "",
        "reviewed_by": "No staff review logged yet",
        "reviewed_at": "",
        "source": "Awaiting public-file review",
    }


def _front_office_store() -> Dict[str, Dict[str, Any]]:
    if "front_office_meta" not in st.session_state:
        st.session_state["front_office_meta"] = {}
    return st.session_state["front_office_meta"]


def _front_office_audit_store() -> Dict[str, List[Dict[str, Any]]]:
    if "front_office_audit" not in st.session_state:
        st.session_state["front_office_audit"] = {}
    return st.session_state["front_office_audit"]


def append_front_office_audit(evaluation_id: str, section: str, actor: str, summary: str) -> None:
    store = _front_office_audit_store()
    history = list(store.get(evaluation_id) or [])
    history.append(
        {
            "at": _now_iso(),
            "section": section,
            "actor": actor,
            "summary": summary,
        }
    )
    store[evaluation_id] = history[-10:]


def get_front_office_audit(evaluation_id: str) -> List[Dict[str, Any]]:
    return list((_front_office_audit_store().get(evaluation_id) or []))


def sync_front_office_meta(evaluations: List[Dict[str, Any]]) -> None:
    store = _front_office_store()
    for row in evaluations:
        eval_id = row.get("id")
        if not eval_id:
            continue
        defaults = derive_front_office_profile(row)
        existing = dict(store.get(eval_id) or {})
        store[eval_id] = defaults | existing


def get_front_office_meta(detail: Dict[str, Any]) -> Dict[str, Any]:
    eval_id = detail.get("id")
    defaults = derive_front_office_profile(detail) | derive_medical_diligence_profile(detail)
    if not eval_id:
        return defaults
    return defaults | dict(_front_office_store().get(eval_id) or {})


def save_front_office_meta(evaluation_id: str, **updates: Any) -> None:
    store = _front_office_store()
    current = dict(store.get(evaluation_id) or {})
    current.update({key: value for key, value in updates.items() if value is not None})
    store[evaluation_id] = current


def push_workspace_notice(section: str, evaluation_id: str, message: str) -> None:
    st.session_state["workspace_notice"] = {
        "section": section,
        "evaluation_id": evaluation_id,
        "message": message,
    }


def render_workspace_notice(section: str, evaluation_id: str) -> None:
    notice = st.session_state.get("workspace_notice")
    if not notice:
        return
    if notice.get("section") != section or notice.get("evaluation_id") != evaluation_id:
        return
    st.success(str(notice.get("message") or "Saved."))
    st.session_state.pop("workspace_notice", None)


def get_latest_audit_entry(evaluation_id: str, section: str) -> Optional[Dict[str, Any]]:
    history = get_front_office_audit(evaluation_id)
    for entry in reversed(history):
        if entry.get("section") == section:
            return entry
    return None

def current_workspace_role() -> str:
    return str(st.session_state.get("workspace_role") or "gm")


def is_collaborator_role() -> bool:
    return current_workspace_role() in COLLABORATOR_ROLES


def is_workspace_authenticated() -> bool:
    return bool(st.session_state.get("workspace_authenticated"))


def current_workspace_user() -> str:
    return str(st.session_state.get("workspace_username") or "")


def render_workspace_login() -> None:
    st.markdown('<div class="section-kicker">Workspace Access</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Sign In</div>', unsafe_allow_html=True)
    st.caption("Use a role-specific login to open the GM, sport science, or medical workspace.")
    if WAIMS_DEMO_MODE:
        st.info("Demo logins: `gm / gm123`, `scientist / sci123`, `medical / med123`")

    with st.form("workspace_login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Sign in")
    if submitted:
        user_key = username.strip().lower()
        config = WORKSPACE_LOGIN_USERS.get(user_key)
        if not config or password != config["password"]:
            st.error("Invalid workspace credentials.")
        else:
            st.session_state["workspace_authenticated"] = True
            st.session_state["workspace_username"] = user_key
            st.session_state["workspace_role"] = config["role"]
            st.success(f"Signed in as {config['label']}.")
            st.rerun()


def _local_score_detail_from_payload(
    payload: Dict[str, Any],
    evaluation_id: Optional[str] = None,
    created_at: Optional[str] = None,
) -> Dict[str, Any]:
    player_dict = dict(payload.get("player", {}) or {})
    ctx_dict = dict(payload.get("ctx", {}) or {})
    mode = payload.get("mode") or "pro_wnba"
    ctx_with_meta = ctx_dict | {"gm_id": "demo-local", "mode": mode}

    player = Player(**player_dict)
    ctx = TeamContext(**ctx_with_meta)
    scorecard = evaluate_single_player(player, ctx)

    return {
        "id": evaluation_id or f"demo-{uuid4().hex[:8]}",
        "gm_id": "demo-local",
        "team_id": ctx_dict.get("team_id"),
        "overall_score": scorecard.overall_score,
        "components": scorecard.components,
        "assumptions": scorecard.assumptions,
        "tension_points": scorecard.tension_points,
        "recommended_action": scorecard.recommended_action,
        "player": player_dict,
        "ctx": ctx_dict,
        "created_at": created_at or _now_iso(),
        "summary_note": payload.get("summary_note"),
        "strengths": payload.get("strengths"),
        "concerns": payload.get("concerns"),
        "mode": mode,
        "display_name": payload.get("display_name"),
    }


def _default_demo_details() -> List[Dict[str, Any]]:
    details: List[Dict[str, Any]] = []
    for index, payload in enumerate(demo_payloads()):
        details.append(
            _local_score_detail_from_payload(
                payload,
                evaluation_id=f"demo-local-{index + 1}",
                created_at=_now_iso(),
            )
        )
    return details


def ensure_local_demo_state() -> None:
    if "local_demo_details" not in st.session_state:
        st.session_state["local_demo_details"] = _default_demo_details()


def _local_demo_details() -> List[Dict[str, Any]]:
    ensure_local_demo_state()
    return list(st.session_state["local_demo_details"])


def _save_local_demo_details(details: List[Dict[str, Any]]) -> None:
    st.session_state["local_demo_details"] = details


def get_evaluations(token: str) -> List[Dict[str, Any]]:
    if WAIMS_DEMO_MODE:
        return [
            {
                "id": row["id"],
                "gm_id": row["gm_id"],
                "team_id": row.get("team_id"),
                "overall_score": row["overall_score"],
                "recommended_action": row["recommended_action"],
                "created_at": row.get("created_at"),
                "player": row.get("player", {}),
                "summary_note": row.get("summary_note"),
                "mode": row.get("mode"),
            }
            for row in _local_demo_details()
        ]
    with httpx.Client(timeout=20) as client:
        r = client.get(f"{API_BASE_URL}/evaluations", headers=api_headers(token))
        r.raise_for_status()
        return r.json()


def get_evaluation_detail(token: str, evaluation_id: str) -> Dict[str, Any]:
    if WAIMS_DEMO_MODE:
        for row in _local_demo_details():
            if row["id"] == evaluation_id:
                return row
        raise KeyError(f"No local demo evaluation found for id {evaluation_id}")
    with httpx.Client(timeout=20) as client:
        r = client.get(
            f"{API_BASE_URL}/evaluations/{evaluation_id}",
            headers=api_headers(token),
        )
        r.raise_for_status()
        return r.json()


def create_evaluation(token: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    if WAIMS_DEMO_MODE:
        details = _local_demo_details()
        new_detail = _local_score_detail_from_payload(payload)
        details.insert(0, new_detail)
        _save_local_demo_details(details)
        return {
            "evaluation_id": new_detail["id"],
            "overall_score": new_detail["overall_score"],
            "components": new_detail["components"],
            "assumptions": new_detail["assumptions"],
            "tension_points": new_detail["tension_points"],
            "recommended_action": new_detail["recommended_action"],
            "player": new_detail["player"],
            "summary_note": new_detail.get("summary_note"),
            "strengths": new_detail.get("strengths"),
            "concerns": new_detail.get("concerns"),
            "mode": new_detail.get("mode"),
        }
    with httpx.Client(timeout=30) as client:
        r = client.post(
            f"{API_BASE_URL}/evaluate-and-save",
            headers=api_headers(token),
            json=payload,
        )
        r.raise_for_status()
        return r.json()


def update_evaluation(token: str, evaluation_id: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    if WAIMS_DEMO_MODE:
        details = _local_demo_details()
        for idx, row in enumerate(details):
            if row["id"] != evaluation_id:
                continue
            updated = _local_score_detail_from_payload(
                payload,
                evaluation_id=evaluation_id,
                created_at=row.get("created_at"),
            )
            updated["gm_id"] = row.get("gm_id", "demo-local")
            details[idx] = updated
            _save_local_demo_details(details)
            return updated
        raise KeyError(f"No local demo evaluation found for id {evaluation_id}")

    with httpx.Client(timeout=30) as client:
        r = client.patch(
            f"{API_BASE_URL}/evaluations/{evaluation_id}",
            headers=api_headers(token),
            json=payload,
        )
        r.raise_for_status()
        return r.json()


def delete_evaluation(token: str, evaluation_id: str) -> Dict[str, Any]:
    if WAIMS_DEMO_MODE:
        details = [row for row in _local_demo_details() if row["id"] != evaluation_id]
        _save_local_demo_details(details)
        return {"ok": True, "deleted_id": evaluation_id}
    with httpx.Client(timeout=20) as client:
        r = client.delete(
            f"{API_BASE_URL}/evaluations/{evaluation_id}",
            headers=api_headers(token),
        )
        r.raise_for_status()
        return r.json()


def format_score(value: Optional[float]) -> str:
    return "—" if value is None else f"{value:.2f}"


def format_dt(value: Optional[str]) -> str:
    if not value:
        return "—"
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
        return dt.strftime("%b %d, %Y • %I:%M %p")
    except Exception:
        return value


def clean_action(action: Optional[str], mode: Optional[str]) -> str:
    raw = (action or "").lower()
    mode = mode or "pro_wnba"
    return ACTION_LABELS.get(mode, ACTION_LABELS["pro_wnba"]).get(raw, raw.title() or "—")


def action_class(action: Optional[str]) -> str:
    action = (action or "").lower()
    if action == "draft":
        return "action-draft"
    if action == "sign":
        return "action-sign"
    if action == "pass":
        return "action-pass"
    return ""


def render_header() -> None:
    env_class = "env-badge-live" if IS_LIVE_ENV else "env-badge-sandbox"
    runtime_label = "Local Demo Mode" if WAIMS_DEMO_MODE else f"API {API_BASE_URL}"
    st.markdown(
        f"""
        <div class="waims-header">
            <div class="waims-kicker">Front Office Briefing Terminal</div>
            <h1 class="waims-title">WAIMS-GM Morning Brief</h1>
            <div class="waims-subtitle">
                Executive review layer for player evaluation, fit, risk, action, and scouting rationale.
            </div>
            <div class="waims-meta-row">
                <div class="env-badge {env_class}">{WAIMS_ENV_LABEL}</div>
                <div class="waims-kicker" style="margin-bottom:0;">{runtime_label}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def get_mode_playbook(mode: str) -> Dict[str, str]:
    return MODE_PLAYBOOKS.get(mode or DEFAULT_MODE, MODE_PLAYBOOKS[DEFAULT_MODE])


def render_mode_focus_banner(mode: str, surface: str, show_label: bool = True) -> None:
    mode = mode or DEFAULT_MODE
    playbook = get_mode_playbook(mode)
    notes = {
        "create": playbook["create_note"],
        "board": playbook["board_note"],
        "dossier": playbook["dossier_note"],
        "compare": playbook["compare_note"],
    }
    surface_label = {
        "create": "Mode Setup",
        "board": "Mode Focus",
        "dossier": "Dossier Lens",
        "compare": "Compare Lens",
    }
    if show_label:
        st.markdown(
            f'<div class="section-kicker" style="margin-top:0.2rem;">{surface_label.get(surface, "Mode Focus")}</div>',
            unsafe_allow_html=True,
        )
    st.markdown(
        f"""
        <div class="soft-card" style="margin-bottom:0.85rem;">
            <div class="mini-label">{MODE_LABELS.get(mode, mode)}</div>
            <div class="board-name" style="font-size:1.05rem;">{playbook['headline']}</div>
            <div class="memo-text" style="margin-top:0.45rem;">
                <strong>Primary question:</strong> {playbook['primary_question']}
                <br/>
                <strong>Weighting bias:</strong> {playbook['weight_note']}
                <br/>
                {notes.get(surface, playbook['dossier_note'])}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def summarize_context(ctx: Dict[str, Any]) -> str:
    return (
        f"Team {ctx.get('team_id', '—')} operating in a {ctx.get('timeline', '—')} window, "
        f"with cap flexibility {ctx.get('cap_flexibility', '—')} and risk tolerance {ctx.get('risk_tolerance', '—')}."
    )


def build_dossier_takeaways(detail: Dict[str, Any]) -> List[str]:
    detail = normalize_detail_for_display(detail)
    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    mode = detail.get("mode") or "pro_wnba"
    components = detail.get("components", {}) or {}
    component_scores = [(label, _component_number(components, key) or 0.0) for key, label in COMPONENT_LABELS.items()]
    component_scores.sort(key=lambda item: item[1], reverse=True)
    strongest_label, strongest_value = component_scores[0]

    needs = ctx.get("needs_by_position", {}) or {}
    pos = player.get("position", "—")
    need_value = needs.get(pos)
    need_line = (
        f"{pos} is a live roster-need position at {need_value:.2f}, which supports the current recommendation."
        if need_value is not None
        else "Roster-need data is limited, so fit should be validated against live board priorities."
    )

    tension_points = detail.get("tension_points", []) or []
    risk_line = tension_points[0] if tension_points else "No major tension points are currently flagged in this file."

    overall_score = format_score(detail.get("overall_score"))
    second_label, second_value = component_scores[1]
    playbook = get_mode_playbook(mode)
    level_delta = derive_level_delta(detail)

    return [
        f"Overall score: {overall_score}. Recommendation: {clean_action(detail.get('recommended_action'), mode)}.",
        f"Top components: {strongest_label} {format_score(strongest_value)} and {second_label} {format_score(second_value)}.",
        f"Bet profile: {level_delta['bet_label']} — {level_delta['level_label']} level with {level_delta['delta_label']} delta.",
        f"Mode lens: {playbook['weight_note']}",
        f"Roster / risk view: {need_line} Primary watch item: {risk_line}",
    ]


def build_executive_brief_cards(detail: Dict[str, Any]) -> List[Dict[str, str]]:
    detail = normalize_detail_for_display(detail)
    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    mode = detail.get("mode") or DEFAULT_MODE
    components = detail.get("components", {}) or {}
    component_scores = [
        (label, _component_number(components, key) or 0.0)
        for key, label in COMPONENT_LABELS.items()
    ]
    component_scores.sort(key=lambda item: item[1], reverse=True)
    strongest_label, strongest_value = component_scores[0]
    second_label, second_value = component_scores[1]

    needs = ctx.get("needs_by_position", {}) or {}
    pos = player.get("position", "—")
    need_value = needs.get(pos)
    tension_points = detail.get("tension_points", []) or []
    primary_watch = tension_points[0] if tension_points else "No major tension points are currently flagged in this file."

    why_now = (
        f"{clean_action(detail.get('recommended_action'), mode)} at {format_score(detail.get('overall_score'))}. "
        f"{strongest_label} ({format_score(strongest_value)}) and {second_label} ({format_score(second_value)}) are driving the file."
    )
    why_not = primary_watch
    decision_driver = (
        f"{pos} need is {need_value:.2f}. {get_mode_playbook(mode)['weight_note']}"
        if need_value is not None
        else get_mode_playbook(mode)["weight_note"]
    )

    return [
        {"title": "Why Now", "winner": clean_action(detail.get("recommended_action"), mode), "note": why_now},
        {"title": "Why Not", "winner": "Primary Watch Item", "note": why_not},
        {"title": "Decision Drivers", "winner": MODE_LABELS.get(mode, mode), "note": decision_driver},
    ]


def render_diagnostic_strip(detail: Dict[str, Any]) -> None:
    rows = compute_five_layer_diagnostic(detail)
    short_labels = {
        "Layer 1 — Scoring Profile": "Scoring",
        "Layer 2 — Creation Profile": "Creation",
        "Layer 3 — Defensive Translation": "Defense",
        "Layer 4 — Availability / Stability": "Stability",
        "Layer 5 — Value / Roster Fit": "Value / Fit",
    }
    st.markdown(
        '<div class="section-kicker" style="margin-top:0.55rem; margin-bottom:0.3rem;">Diagnostic Snapshot</div>',
        unsafe_allow_html=True,
    )
    cols = st.columns(len(rows), gap="small")
    for col, row in zip(cols, rows):
        with col:
            st.markdown(
                f"""
                <div class="diagnostic-mini-card">
                    <div class="diagnostic-mini-label">{short_labels.get(row['layer'], row['layer'])}</div>
                    <div class="diagnostic-mini-grade">{row['grade']}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )


def render_level_delta_note() -> None:
    st.caption(
        "Level / Delta is a simple bet-type lens: Level = expected contribution band, Delta = how much the outcome could swing. "
        "Credit: John Chisholm. Reference: Rebelo, et al. (2026), Monitoring Training Effects in Athletes: A Multidimensional Framework for Decision-Making, Sports Medicine."
    )


def _player_rating_0_100(player: Dict[str, Any], key: str) -> float:
    try:
        value = float(player.get(key, 0) or 0)
    except (TypeError, ValueError):
        return 0.0
    if 0.0 <= value <= 1.0:
        value *= 100.0
    return max(0.0, min(value, 100.0))


def compute_archetype_mix(detail: Dict[str, Any], top_n: int = 3) -> List[Dict[str, float]]:
    player = detail.get("player", {}) or {}
    position = str(player.get("position", "") or "").upper()

    offense = _player_rating_0_100(player, "offense_rating")
    defense = _player_rating_0_100(player, "defense_rating")
    shooting = _player_rating_0_100(player, "shooting_rating")
    playmaking = _player_rating_0_100(player, "playmaking_rating")
    rebounding = _player_rating_0_100(player, "rebounding_rating")

    archetype_scores = {
        "Lead Creator": playmaking * 0.42 + offense * 0.28 + shooting * 0.18 + defense * 0.07 + rebounding * 0.05,
        "Secondary Creator": playmaking * 0.28 + shooting * 0.24 + offense * 0.22 + defense * 0.16 + rebounding * 0.10,
        "Floor Spacer": shooting * 0.52 + offense * 0.20 + defense * 0.12 + playmaking * 0.10 + rebounding * 0.06,
        "3-and-D Wing": defense * 0.34 + shooting * 0.30 + offense * 0.16 + rebounding * 0.12 + playmaking * 0.08,
        "Connector Forward": playmaking * 0.24 + defense * 0.24 + rebounding * 0.20 + offense * 0.16 + shooting * 0.16,
        "Interior Big": rebounding * 0.34 + defense * 0.28 + offense * 0.22 + shooting * 0.08 + playmaking * 0.08,
    }

    position_bonus = {
        "Lead Creator": {"G": 6.0, "F": 1.0, "C": -8.0},
        "Secondary Creator": {"G": 4.0, "F": 2.0, "C": -6.0},
        "Floor Spacer": {"G": 3.0, "F": 3.0, "C": -4.0},
        "3-and-D Wing": {"G": 3.0, "F": 6.0, "C": -6.0},
        "Connector Forward": {"G": -2.0, "F": 6.0, "C": 1.0},
        "Interior Big": {"G": -10.0, "F": 2.0, "C": 8.0},
    }

    scored = []
    for label, score in archetype_scores.items():
        bonus = position_bonus.get(label, {}).get(position, 0.0)
        scored.append({"label": label, "score": max(0.0, score + bonus)})

    scored.sort(key=lambda item: item["score"], reverse=True)
    selected = scored[: max(1, top_n)]
    total = sum(item["score"] for item in selected) or 1.0

    return [
        {
            "label": item["label"],
            "score": round(item["score"], 1),
            "share": round((item["score"] / total) * 100.0, 1),
        }
        for item in selected
    ]


def summarize_archetype_mix(detail: Dict[str, Any], top_n: int = 3) -> str:
    mix = compute_archetype_mix(detail, top_n=top_n)
    if not mix:
        return "No archetype mix available."
    return ", ".join(f"{item['label']} ({item['share']:.0f}%)" for item in mix)


def build_archetype_profile_card(detail: Dict[str, Any]) -> Dict[str, str]:
    mix = compute_archetype_mix(detail)
    if not mix:
        return {
            "title": "Archetype Mix",
            "winner": "Not available",
            "note": "No role mix is available for this file yet.",
        }

    primary = mix[0]["label"]
    note = f"Role mix: {summarize_archetype_mix(detail)}."
    if len(mix) > 1:
        note += " Use this as a style-fit layer on top of the current score and Five Layers diagnostic."

    return {
        "title": "Archetype Mix",
        "winner": primary,
        "note": note,
    }


def build_archetype_comparison_card(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> Dict[str, str]:
    left_name = (left_detail.get("player") or {}).get("name", "Selected Player")
    right_name = (right_detail.get("player") or {}).get("name", "Comparison Player")
    left_mix = compute_archetype_mix(left_detail, top_n=2)
    right_mix = compute_archetype_mix(right_detail, top_n=2)

    left_labels = [item["label"] for item in left_mix]
    right_labels = [item["label"] for item in right_mix]
    shared = [label for label in left_labels if label in right_labels]

    complementary_pairs = {
        ("Lead Creator", "Floor Spacer"),
        ("Lead Creator", "3-and-D Wing"),
        ("Lead Creator", "Connector Forward"),
        ("Secondary Creator", "3-and-D Wing"),
        ("Connector Forward", "Interior Big"),
        ("Floor Spacer", "Interior Big"),
        ("3-and-D Wing", "Interior Big"),
    }
    top_pair = tuple(sorted([left_labels[0], right_labels[0]])) if left_labels and right_labels else ()

    if len(shared) >= 2 or (left_labels and right_labels and left_labels[0] == right_labels[0]):
        winner = "High overlap"
        note = (
            f"Both files lean toward the same role family. {left_name} reads as {summarize_archetype_mix(left_detail, 2)}, "
            f"while {right_name} reads as {summarize_archetype_mix(right_detail, 2)}. This is useful if you want role insurance, "
            "but it creates more duplication risk."
        )
    elif top_pair in complementary_pairs:
        winner = "Complementary"
        note = (
            f"The styles can fit together cleanly: {left_name} reads as {summarize_archetype_mix(left_detail, 2)}, "
            f"while {right_name} reads as {summarize_archetype_mix(right_detail, 2)}. This points more toward lineup complement than role duplication."
        )
    elif shared:
        winner = "Some overlap"
        note = (
            f"The files share some stylistic overlap, but not enough to be interchangeable. "
            f"{left_name}: {summarize_archetype_mix(left_detail, 2)}. {right_name}: {summarize_archetype_mix(right_detail, 2)}."
        )
    else:
        winner = "Different styles"
        note = (
            f"The role identities are meaningfully different. {left_name}: {summarize_archetype_mix(left_detail, 2)}. "
            f"{right_name}: {summarize_archetype_mix(right_detail, 2)}. This is a cleaner fit conversation than a pure score tie-break."
        )

    return {"title": "Archetype Fit", "winner": winner, "note": note}


def build_memo(detail: Dict[str, Any]) -> str:
    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    summary_note = detail.get("summary_note") or ""
    base = (
        f"{player.get('name', 'This player')} profiles as a {player.get('position', '—')} target with an overall "
        f"score of {format_score(detail.get('overall_score'))}. The current recommendation is "
        f"{clean_action(detail.get('recommended_action'), detail.get('mode'))}. Age ({player.get('age', '—')}) and expected cost tier "
        f"({player.get('expected_cost_tier', '—')}) frame the acquisition lens, while the evaluation context points "
        f"toward {ctx.get('timeline', 'a flexible')} team-building priorities rather than a purely static rating interpretation."
    )
    if summary_note:
        base += f" Analyst note: {summary_note}"
    return base


def prepare_evaluations(
    evaluations: List[Dict[str, Any]],
    action_filter: str,
    hide_placeholder: bool,
    mode_filter: str,
    stage_filter: str,
) -> List[Dict[str, Any]]:
    rows = list(evaluations)
    if hide_placeholder:
        rows = [
            r for r in rows
            if (r.get("player", {}) or {}).get("name", "").strip().lower() not in {"string", ""}
        ]
    if action_filter != "All":
        rows = [r for r in rows if (r.get("recommended_action") or "").lower() == action_filter.lower()]
    if mode_filter != "All":
        rows = [r for r in rows if (r.get("mode") or "pro_wnba") == mode_filter]
    if stage_filter != "All":
        rows = [r for r in rows if get_front_office_meta(r).get("stage") == stage_filter]
    return rows


def sort_evaluations(rows: List[Dict[str, Any]], sort_by: str, descending: bool) -> List[Dict[str, Any]]:
    def key_fn(row: Dict[str, Any]):
        if sort_by == "Score":
            return float(row.get("overall_score") or 0)
        if sort_by == "Recommendation":
            return str(row.get("recommended_action") or "")
        if sort_by == "Mode":
            return str(row.get("mode") or "")
        if sort_by == "Player Name":
            return str((row.get("player") or {}).get("name") or "")
        return str(row.get("created_at") or "")
    return sorted(rows, key=key_fn, reverse=descending)


def render_summary_cards(evaluations: List[Dict[str, Any]]) -> None:
    total = len(evaluations)
    draft_count = sum(1 for e in evaluations if e.get("recommended_action") == "draft")
    sign_count = sum(1 for e in evaluations if e.get("recommended_action") == "sign")
    priority_count = sum(1 for e in evaluations if get_front_office_meta(e).get("stage") in {"Priority", "Live Board"})
    avg_score = (
        sum(float(e.get("overall_score", 0)) for e in evaluations) / total if total else 0
    )
    safe_rotation_count = sum(1 for e in evaluations if derive_level_delta(e).get("bet_label") == "Safe Rotation Bet")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(
            f"""<div class="metric-card"><div class="metric-label">Saved Evaluations</div><div class="metric-value">{total}</div><div class="metric-note">Current briefing inventory</div></div>""",
            unsafe_allow_html=True,
        )
    with col2:
        st.markdown(
            f"""<div class="metric-card"><div class="metric-label">Average Score</div><div class="metric-value">{avg_score:.2f}</div><div class="metric-note">Across filtered evaluations</div></div>""",
            unsafe_allow_html=True,
        )
    with col3:
        st.markdown(
            f"""<div class="metric-card"><div class="metric-label">Priority Queue</div><div class="metric-value">{priority_count}</div><div class="metric-note">{draft_count} draft / {sign_count} sign recommendations still on board</div></div>""",
            unsafe_allow_html=True,
        )
    with col4:
        st.markdown(
            f"""<div class="metric-card"><div class="metric-label">Safe Rotation Bets</div><div class="metric-value">{safe_rotation_count}</div><div class="metric-note">Reliable middle-tier contributors on the current board</div></div>""",
            unsafe_allow_html=True,
        )


def render_empty_board_state() -> None:
    st.markdown(
        """
        <div class="soft-card" style="margin-top:0.85rem;">
            <div class="mini-label">Board Status</div>
            <div class="board-name" style="font-size:1.05rem;">No saved evaluations match the current view.</div>
            <div class="memo-text" style="margin-top:0.45rem;">
                Start by creating a new evaluation on the left, clear the active filters, or seed repeatable demo files with
                <code>scripts/seed_demo_data.py</code>.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_decision_board(evaluations: List[Dict[str, Any]]) -> Optional[str]:
    st.markdown('<div class="section-kicker">Priority Queue</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Recent Evaluation Board</div>', unsafe_allow_html=True)

    if not evaluations:
        render_empty_board_state()
        return None

    existing_ids = {row["id"] for row in evaluations}
    selected_id = st.session_state.get("selected_evaluation_id")
    if selected_id not in existing_ids:
        selected_id = evaluations[0]["id"]
        st.session_state["selected_evaluation_id"] = selected_id

    for item in evaluations[:12]:
        player = item.get("player", {}) or {}
        note = item.get("summary_note") or ""
        selected = item["id"] == selected_id
        extra_class = " board-card-selected" if selected else ""
        mode = item.get("mode") or "pro_wnba"
        workflow = get_front_office_meta(item)
        workflow_line = (
            f"{workflow.get('stage', 'Shadow Board')} &nbsp;|&nbsp; "
            f"{workflow.get('priority_tier', 'Tier 3')} &nbsp;|&nbsp; "
            f"Owner {workflow.get('owner', 'GM')}"
        )
        valuation_line = (
            f"{workflow.get('value_posture', 'Price-sensitive')} &nbsp;|&nbsp; "
            f"{workflow.get('market_band', 'Market')} spend band"
        )
        level_delta = derive_level_delta(item)
        bet_line = (
            f"Level / Delta: {level_delta.get('bet_label')} &nbsp;|&nbsp; "
            f"{level_delta.get('level_label')} &nbsp;|&nbsp; {level_delta.get('delta_label')}"
        )
        diligence_line = (
            f"{MEDICAL_DILIGENCE_LABEL}: {workflow.get('level', 'Proceed with caution')}"
            f" &nbsp;|&nbsp; Confidence {workflow.get('confidence', 'Medium')}"
        )

        st.markdown(
            f"""
            <div class="board-card{extra_class}">
                <div class="board-head">
                    <div class="board-name">{player.get("name", "Unknown Player")}</div>
                    <div class="board-tag {action_class(item.get("recommended_action"))}">{clean_action(item.get("recommended_action"), mode)}</div>
                </div>
                <div class="board-meta">
                    {MODE_LABELS.get(mode, mode)} &nbsp;|&nbsp; {player.get("position", "—")}
                    &nbsp;|&nbsp; Score {format_score(item.get("overall_score"))}
                    &nbsp;|&nbsp; {format_dt(item.get("created_at"))}
                </div>
                <div class="board-meta">{workflow_line}</div>
                <div class="board-meta">{valuation_line}</div>
                <div class="board-meta">{bet_line}</div>
                <div class="board-meta">{diligence_line}</div>
                {"<div class='board-note'>" + note + "</div>" if note else ""}
                {"<div class='board-note'><strong>Next:</strong> " + str(workflow.get("next_action") or "") + "</div>" if workflow.get("next_action") else ""}
            </div>
            """,
            unsafe_allow_html=True,
        )

        if st.button("Selected" if selected else f"Open {player.get('name', 'Player')}", key=f"open_{item['id']}", disabled=selected):
            st.session_state["selected_evaluation_id"] = item["id"]
            st.rerun()

    return st.session_state.get("selected_evaluation_id")


def render_bullets(text_block: Optional[str], fallback: str) -> str:
    if not text_block or not text_block.strip():
        return f"<li>{fallback}</li>"
    lines = [line.strip("•- ").strip() for line in text_block.splitlines() if line.strip()]
    if not lines:
        return f"<li>{fallback}</li>"
    return "".join(f"<li>{line}</li>" for line in lines)


def text_block_lines(text_block: Optional[str], fallback: Optional[str] = None) -> List[str]:
    lines = [line.strip("•- ").strip() for line in (text_block or "").splitlines() if line.strip()]
    if lines:
        return lines
    return [fallback] if fallback else []


def normalize_detail_for_display(detail: Dict[str, Any]) -> Dict[str, Any]:
    normalized = dict(detail)
    raw_components = dict(normalized.get("components") or {})
    components: Dict[str, Optional[float]] = {}

    for key in COMPONENT_LABELS:
        value = raw_components.get(key)
        try:
            components[key] = None if value in (None, "") else float(value)
        except (TypeError, ValueError):
            components[key] = None

    legacy_availability = raw_components.get("risk_inverse")
    try:
        if components["availability"] is None and legacy_availability not in (None, ""):
            legacy_value = float(legacy_availability)
            components["availability"] = legacy_value * 100 if legacy_value <= 1 else legacy_value
    except (TypeError, ValueError):
        pass

    if any(value is None for value in components.values()):
        player_payload = normalized.get("player") or {}
        ctx_payload = dict(normalized.get("ctx") or {})

        if player_payload and ctx_payload:
            try:
                ctx_payload.setdefault("gm_id", str(normalized.get("gm_id") or "legacy-display"))
                ctx_payload.setdefault("mode", normalized.get("mode") or ctx_payload.get("mode") or "pro_wnba")
                rescored = evaluate_single_player(
                    Player(**player_payload),
                    TeamContext(**ctx_payload),
                )
                for key, value in rescored.components.items():
                    if components.get(key) is None:
                        components[key] = float(value)
            except Exception:
                pass

    normalized["components"] = {
        key: round(value, 2)
        for key, value in components.items()
        if value is not None
    }
    return normalized


def _component_number(components: Dict[str, Any], key: str) -> Optional[float]:
    value = components.get(key)
    try:
        return None if value in (None, "") else float(value)
    except (TypeError, ValueError):
        return None


def build_decision_lens(detail: Dict[str, Any]) -> str:
    normalized = normalize_detail_for_display(detail)
    components = normalized.get("components", {}) or {}
    mode = normalized.get("mode") or "pro_wnba"

    metric_blurbs = {
        "fit": "current roster fit",
        "impact": "immediate impact",
        "upside": "long-term upside",
        "availability": "near-term dependability",
        "value": "value-to-cost profile",
    }
    mode_lenses = {
        "pro_wnba": "win-now acquisition",
        "cbb_high_major": "high-major translation",
        "cbb_d2_low_resource": "small-staff roster",
        "recruiting_only": "long-horizon recruiting",
    }

    ranked = [
        (key, value)
        for key in COMPONENT_LABELS
        if (value := _component_number(components, key)) is not None
    ]
    if not ranked:
        return "This file still needs more complete component data before a sharper acquisition lens can be stated."

    ranked.sort(key=lambda item: item[1], reverse=True)
    strongest_key, strongest_value = ranked[0]
    weakest_key, weakest_value = ranked[-1]
    recommendation = clean_action(normalized.get("recommended_action"), mode)

    return (
        f"This profile reads most strongly through a {mode_lenses.get(mode, 'basketball decision')} lens: "
        f"{metric_blurbs[strongest_key]} is the main reason the file carries a {recommendation} recommendation "
        f"({strongest_value:.1f}), while {metric_blurbs[weakest_key]} remains the cleanest pressure point "
        f"to monitor ({weakest_value:.1f})."
    )


def derive_level_delta(detail: Dict[str, Any]) -> Dict[str, str]:
    normalized = normalize_detail_for_display(detail)
    components = normalized.get("components", {}) or {}
    player = normalized.get("player", {}) or {}
    mode = normalized.get("mode") or DEFAULT_MODE

    fit = _component_number(components, "fit") or 0.0
    impact = _component_number(components, "impact") or 0.0
    availability = _component_number(components, "availability") or 0.0
    value = _component_number(components, "value") or 0.0
    upside_component = _component_number(components, "upside") or 0.0
    health_risk = float(player.get("health_risk", 0) or 0)
    minutes_stability = float(player.get("minutes_stability", 0) or 0)

    level_weights = {
        "pro_wnba": {"fit": 0.28, "impact": 0.36, "availability": 0.22, "value": 0.14},
        "cbb_high_major": {"fit": 0.30, "impact": 0.30, "availability": 0.18, "value": 0.12},
        "cbb_d2_low_resource": {"fit": 0.32, "impact": 0.18, "availability": 0.24, "value": 0.26},
        "recruiting_only": {"fit": 0.18, "impact": 0.16, "availability": 0.12, "value": 0.12},
    }.get(mode, {"fit": 0.30, "impact": 0.30, "availability": 0.20, "value": 0.20})
    level_score = (
        fit * level_weights["fit"]
        + impact * level_weights["impact"]
        + availability * level_weights["availability"]
        + value * level_weights["value"]
    )

    delta_score = (
        upside_component * 0.45
        + max(0.0, (health_risk * 100)) * 0.25
        + max(0.0, (100 - minutes_stability * 100)) * 0.30
    )

    if level_score >= 80:
        level_label = "Impact Player"
        level_note = "Projects more like a real lineup driver than a depth-only file if the plan holds."
    elif level_score >= 68:
        level_label = "Rotation Value"
        level_note = "Profiles as a usable rotation contributor, which is where smaller staffs create a lot of value."
    else:
        level_label = "Replacement Risk"
        level_note = "Needs more things to break right before this reads like a reliable contributor bet."

    if delta_score >= 62:
        delta_label = "High Swing"
        delta_note = "Big outcome range: the file carries upside, but also more volatility and miss risk."
    elif delta_score >= 42:
        delta_label = "Moderate Swing"
        delta_note = "Some upside and some uncertainty, but not an all-or-nothing profile."
    else:
        delta_label = "Low Swing"
        delta_note = "Cleaner, steadier bet. The likely outcome band is tighter and easier to price."

    bet_label_map = {
        ("Impact Player", "Low Swing"): "Safe Impact Bet",
        ("Impact Player", "Moderate Swing"): "High-Level Bet",
        ("Impact Player", "High Swing"): "Star Swing",
        ("Rotation Value", "Low Swing"): "Safe Rotation Bet",
        ("Rotation Value", "Moderate Swing"): "Rotation Bet",
        ("Rotation Value", "High Swing"): "Upside Rotation Swing",
        ("Replacement Risk", "Low Swing"): "Stable Depth Add",
        ("Replacement Risk", "Moderate Swing"): "Thin-Margin Bet",
        ("Replacement Risk", "High Swing"): "Risky File",
    }
    bet_label = bet_label_map.get((level_label, delta_label), "Mixed Bet")
    bet_note = f"{level_note} {delta_note}"

    return {
        "level_score": f"{level_score:.1f}",
        "delta_score": f"{delta_score:.1f}",
        "level_label": level_label,
        "delta_label": delta_label,
        "bet_label": bet_label,
        "bet_note": bet_note,
    }


def build_level_delta_cards(detail: Dict[str, Any]) -> List[Dict[str, str]]:
    lens = derive_level_delta(detail)
    return [
        {
            "title": "Level",
            "winner": lens["level_label"],
            "note": f"Expected return band ({lens['level_score']}).",
        },
        {
            "title": "Delta",
            "winner": lens["delta_label"],
            "note": f"Outcome swing / variance ({lens['delta_score']}).",
        },
        {
            "title": "Bet Type",
            "winner": lens["bet_label"],
            "note": lens["bet_note"],
        },
    ]


def render_level_delta_section(detail: Dict[str, Any]) -> None:
    st.markdown('<div class="section-kicker" style="margin-top:0.4rem;">Bet-Type Lens</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Level / Delta</div>', unsafe_allow_html=True)
    render_soft_card_grid(build_level_delta_cards(detail), columns_per_row=3, top_margin="0")
    render_level_delta_note()


def build_comparison_verdicts(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> List[Dict[str, str]]:
    left_name = (left_detail.get("player") or {}).get("name", "Selected Player")
    right_name = (right_detail.get("player") or {}).get("name", "Comparison Player")
    left_components = left_detail.get("components", {}) or {}
    right_components = right_detail.get("components", {}) or {}
    mode = left_detail.get("mode") or right_detail.get("mode") or "pro_wnba"
    left_lens = derive_level_delta(left_detail)
    right_lens = derive_level_delta(right_detail)

    label_map = {
        "pro_wnba": "Best current roster target",
        "cbb_high_major": "Best high-major fit",
        "cbb_d2_low_resource": "Best current roster target",
        "recruiting_only": "Best long-range target",
    }
    verdict_specs = [
        (label_map.get(mode, "Best current roster target"), "fit", "fit"),
        ("Safer near-term bet", "availability", "availability"),
        ("Higher long-term upside", "upside", "upside"),
        ("Better value profile", "value", "value"),
    ]

    verdicts: List[Dict[str, str]] = []
    for title, key, blurb in verdict_specs:
        left_value = _component_number(left_components, key)
        right_value = _component_number(right_components, key)
        if left_value is None and right_value is None:
            continue

        if left_value is None:
            winner = right_name
            margin = None
        elif right_value is None:
            winner = left_name
            margin = None
        elif abs(left_value - right_value) < 0.01:
            winner = "Even"
            margin = 0.0
        elif left_value > right_value:
            winner = left_name
            margin = left_value - right_value
        else:
            winner = right_name
            margin = right_value - left_value

        if winner == "Even":
            note = f"Both players grade essentially even on {blurb}."
        elif margin is None:
            note = f"{winner} is the only file with usable {blurb} data in this comparison."
        else:
            note = f"{winner} leads {blurb} by {margin:.1f} points."

        verdicts.append({"title": title, "winner": winner, "note": note})

    verdicts.append(
        {
            "title": "Bet Profile",
            "winner": left_name if left_lens["delta_label"] == "Low Swing" and right_lens["delta_label"] != "Low Swing" else (
                right_name if right_lens["delta_label"] == "Low Swing" and left_lens["delta_label"] != "Low Swing" else "Context call"
            ),
            "note": f"{left_name}: {left_lens['bet_label']} | {right_name}: {right_lens['bet_label']}.",
        }
    )

    return verdicts


def build_compare_decision_snapshot(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> List[Dict[str, str]]:
    left_name = (left_detail.get("player") or {}).get("name", "Selected Player")
    right_name = (right_detail.get("player") or {}).get("name", "Comparison Player")
    mode = left_detail.get("mode") or right_detail.get("mode") or "pro_wnba"

    lane_weights = {
        "pro_wnba": {
            "win_now": {"fit": 0.34, "impact": 0.34, "availability": 0.22, "value": 0.07, "upside": 0.03},
            "asset": {"fit": 0.18, "impact": 0.18, "availability": 0.14, "value": 0.18, "upside": 0.32},
            "efficiency": {"fit": 0.22, "impact": 0.12, "availability": 0.16, "value": 0.40, "upside": 0.10},
        },
        "cbb_high_major": {
            "win_now": {"fit": 0.32, "impact": 0.31, "availability": 0.16, "value": 0.07, "upside": 0.14},
            "asset": {"fit": 0.20, "impact": 0.13, "availability": 0.10, "value": 0.17, "upside": 0.40},
            "efficiency": {"fit": 0.25, "impact": 0.10, "availability": 0.12, "value": 0.38, "upside": 0.15},
        },
        "cbb_d2_low_resource": {
            "win_now": {"fit": 0.32, "impact": 0.18, "availability": 0.22, "value": 0.22, "upside": 0.06},
            "asset": {"fit": 0.18, "impact": 0.08, "availability": 0.10, "value": 0.19, "upside": 0.45},
            "efficiency": {"fit": 0.26, "impact": 0.08, "availability": 0.16, "value": 0.42, "upside": 0.08},
        },
        "recruiting_only": {
            "win_now": {"fit": 0.22, "impact": 0.10, "availability": 0.10, "value": 0.12, "upside": 0.46},
            "asset": {"fit": 0.18, "impact": 0.06, "availability": 0.06, "value": 0.12, "upside": 0.58},
            "efficiency": {"fit": 0.21, "impact": 0.05, "availability": 0.08, "value": 0.28, "upside": 0.38},
        },
    }
    lane_meta = {
        "win_now": ("Best current rotation answer", "best matches the immediate roster problem"),
        "asset": ("Best long-term asset bet", "carries the strongest future-facing return profile"),
        "efficiency": ("Best value decision", "gives the cleanest return relative to cost and risk"),
    }
    weights_by_lane = lane_weights.get(mode, lane_weights["pro_wnba"])

    def lane_score(detail: Dict[str, Any], weights: Dict[str, float]) -> float:
        components = detail.get("components", {}) or {}
        return sum((_component_number(components, key) or 0.0) * weight for key, weight in weights.items())

    snapshot = []
    for lane, weights in weights_by_lane.items():
        left_score = lane_score(left_detail, weights)
        right_score = lane_score(right_detail, weights)
        title, rationale = lane_meta[lane]

        if abs(left_score - right_score) < 0.5:
            winner = "Even"
            note = f"Both files land nearly level here, so budget and staff conviction should break the tie."
        elif left_score > right_score:
            winner = left_name
            note = f"{left_name} {rationale} by {left_score - right_score:.1f} points."
        else:
            winner = right_name
            note = f"{right_name} {rationale} by {right_score - left_score:.1f} points."

        snapshot.append({"title": title, "winner": winner, "note": note})

    return snapshot


def build_roster_need_call(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> Dict[str, str]:
    left_name = (left_detail.get("player") or {}).get("name", "Selected Player")
    right_name = (right_detail.get("player") or {}).get("name", "Comparison Player")
    left_player = left_detail.get("player", {}) or {}
    right_player = right_detail.get("player", {}) or {}
    ctx = left_detail.get("ctx") or right_detail.get("ctx") or {}
    mode = left_detail.get("mode") or right_detail.get("mode") or "pro_wnba"

    mode_titles = {
        "pro_wnba": "Roster Need Call",
        "cbb_high_major": "High-Major Need Call",
        "cbb_d2_low_resource": "Roster Need Call",
        "recruiting_only": "Recruiting Need Call",
    }
    mode_lenses = {
        "pro_wnba": "This lens leans toward fit, impact, and dependable availability.",
        "cbb_high_major": "This lens leans toward fit, impact, and high-major translation.",
        "cbb_d2_low_resource": "This lens leans toward fit, value, and near-term reliability.",
        "recruiting_only": "This lens leans toward upside, fit, and long-horizon value.",
    }
    compare_weights = {
        "pro_wnba": {"fit": 0.34, "impact": 0.30, "availability": 0.20, "value": 0.10, "upside": 0.06},
        "cbb_high_major": {"fit": 0.32, "impact": 0.28, "availability": 0.14, "value": 0.10, "upside": 0.16},
        "cbb_d2_low_resource": {"fit": 0.30, "impact": 0.12, "availability": 0.18, "value": 0.30, "upside": 0.10},
        "recruiting_only": {"fit": 0.22, "impact": 0.08, "availability": 0.05, "value": 0.15, "upside": 0.50},
    }
    weights = compare_weights.get(mode, compare_weights["pro_wnba"])
    needs = ctx.get("needs_by_position", {}) or {}

    def _position_need(player: Dict[str, Any]) -> float:
        try:
            return float(needs.get(player.get("position"), 0.5)) * 100
        except (TypeError, ValueError):
            return 50.0

    def _blended_score(detail: Dict[str, Any]) -> float:
        components = detail.get("components", {}) or {}
        component_score = sum(
            (_component_number(components, key) or 0.0) * weight
            for key, weight in weights.items()
        )
        return component_score * 0.78 + _position_need(detail.get("player", {}) or {}) * 0.22

    left_score = _blended_score(left_detail)
    right_score = _blended_score(right_detail)

    if abs(left_score - right_score) < 0.5:
        winner = "Even"
        margin_note = "The files land close enough that the tie should be broken by budget, staff preference, or live intel."
    elif left_score > right_score:
        winner = left_name
        margin_note = f"{left_name} fits the current need stack more cleanly by {left_score - right_score:.1f} points."
    else:
        winner = right_name
        margin_note = f"{right_name} fits the current need stack more cleanly by {right_score - left_score:.1f} points."

    primary_need = None
    if needs:
        try:
            primary_need = max(needs.items(), key=lambda item: float(item[1]))[0]
        except (TypeError, ValueError):
            primary_need = None

    left_pos = left_player.get("position", "—")
    right_pos = right_player.get("position", "—")
    need_note = (
        f"The sharpest current roster need is {primary_need}, with {left_name} profiled at {left_pos} and {right_name} at {right_pos}."
        if primary_need
        else f"{left_name} profiles at {left_pos} and {right_name} at {right_pos}, so the call is being made mostly on component shape rather than a strong positional need."
    )

    return {
        "title": mode_titles.get(mode, "Roster Need Call"),
        "winner": winner,
        "note": f"{margin_note} {need_note} {mode_lenses.get(mode, '')}".strip(),
    }


def compute_five_layer_diagnostic(detail: Dict[str, Any]) -> List[Dict[str, str]]:
    player = detail.get("player", {}) or {}
    offense = float(player.get("offense_rating", 0))
    defense = float(player.get("defense_rating", 0))
    shooting = float(player.get("shooting_rating", 0))
    playmaking = float(player.get("playmaking_rating", 0))
    rebounding = float(player.get("rebounding_rating", 0))
    health_risk = float(player.get("health_risk", 0))
    upside = float(player.get("upside", 0))
    minutes = float(player.get("minutes_stability", 0))
    cost = float(player.get("expected_cost_tier", 0))
    overall = float(detail.get("overall_score", 0))

    def grade(v: float) -> str:
        if v >= 80:
            return "A"
        if v >= 70:
            return "B"
        if v >= 60:
            return "C"
        return "D"

    scoring_val = (offense + shooting) / 2
    creation_val = (playmaking + offense) / 2
    defense_val = (defense + rebounding * 0.3) / 1.3
    availability_val = ((1 - health_risk) * 100 * 0.5) + (minutes * 100 * 0.5)
    value_fit_val = min(100, overall * 3.5 + (100 - cost * 10) * 0.25 + upside * 25)

    return [
        {"layer": "Layer 1 — Scoring Profile", "grade": grade(scoring_val), "note": f"Shooting/offense blend indicates a {grade(scoring_val)} scoring profile for current target context."},
        {"layer": "Layer 2 — Creation Profile", "grade": grade(creation_val), "note": f"Playmaking plus offensive creation suggest {'primary' if creation_val >= 75 else 'secondary'} initiator utility."},
        {"layer": "Layer 3 — Defensive Translation", "grade": grade(defense_val), "note": f"Defensive rating and support traits point to a {grade(defense_val)} defensive portability band."},
        {"layer": "Layer 4 — Availability / Stability", "grade": grade(availability_val), "note": f"Health risk and minutes stability combine into a {grade(availability_val)} dependability signal."},
        {"layer": "Layer 5 — Value / Roster Fit", "grade": grade(value_fit_val), "note": f"Overall score, cost tier, upside, and team context combine to a {grade(value_fit_val)} acquisition value fit."},
    ]


def compare_summary(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> str:
    left_name = (left_detail.get("player") or {}).get("name", "Selected Player")
    right_name = (right_detail.get("player") or {}).get("name", "Comparison Player")
    left_score = float(left_detail.get("overall_score") or 0)
    right_score = float(right_detail.get("overall_score") or 0)

    if left_score > right_score:
        winner = left_name
        margin = left_score - right_score
    elif right_score > left_score:
        winner = right_name
        margin = right_score - left_score
    else:
        winner = "Neither player"
        margin = 0.0

    left_player = left_detail.get("player") or {}
    right_player = right_detail.get("player") or {}

    left_strength = max(
        [
            ("offense", float(left_player.get("offense_rating", 0))),
            ("defense", float(left_player.get("defense_rating", 0))),
            ("shooting", float(left_player.get("shooting_rating", 0))),
            ("playmaking", float(left_player.get("playmaking_rating", 0))),
            ("rebounding", float(left_player.get("rebounding_rating", 0))),
        ],
        key=lambda x: x[1],
    )[0]

    right_strength = max(
        [
            ("offense", float(right_player.get("offense_rating", 0))),
            ("defense", float(right_player.get("defense_rating", 0))),
            ("shooting", float(right_player.get("shooting_rating", 0))),
            ("playmaking", float(right_player.get("playmaking_rating", 0))),
            ("rebounding", float(right_player.get("rebounding_rating", 0))),
        ],
        key=lambda x: x[1],
    )[0]

    left_lens = derive_level_delta(left_detail)
    right_lens = derive_level_delta(right_detail)

    if winner == "Neither player":
        return (
            f"{left_name} and {right_name} are effectively level on overall score. "
            f"{left_name} leans most heavily on {left_strength}, while {right_name} leans most heavily on {right_strength}. "
            f"{left_name} reads like a {left_lens['bet_label']}, while {right_name} reads like a {right_lens['bet_label']}. "
            "The better choice depends on roster context, cost sensitivity, and role need."
        )

    return (
        f"{winner} grades better overall by {margin:.2f} points. "
        f"{left_name}'s strongest profile area is {left_strength}, while {right_name}'s strongest profile area is {right_strength}. "
        f"{left_name} profiles as a {left_lens['bet_label']}, while {right_name} profiles as a {right_lens['bet_label']}. "
        "The decision should now be framed through role fit, price discipline, and timeline rather than raw score alone."
    )


def render_score_cards(detail: Dict[str, Any], components: Dict[str, Any]) -> None:
    recommendation = clean_action(detail.get("recommended_action"), detail.get("mode"))

    def fmt(v: Any) -> str:
        if v in (None, "", "—"):
            return "—"
        try:
            return f"{float(v):.1f}"
        except Exception:
            return str(v)

    st.markdown(
        f"""
        <div class="score-grid">
            <div class="score-card">
                <div class="score-label">Fit</div>
                <div class="score-value">{fmt(components.get("fit"))}</div>
            </div>
            <div class="score-card">
                <div class="score-label">Impact</div>
                <div class="score-value">{fmt(components.get("impact"))}</div>
            </div>
            <div class="score-card">
                <div class="score-label">Upside</div>
                <div class="score-value">{fmt(components.get("upside"))}</div>
            </div>
            <div class="score-card">
                <div class="score-label">Availability</div>
                <div class="score-value">{fmt(components.get("availability"))}</div>
            </div>
            <div class="score-card">
                <div class="score-label">Value</div>
                <div class="score-value">{fmt(components.get("value"))}</div>
            </div>
            <div class="score-card score-card-wide">
                <div class="score-label">Recommendation</div>
                <div class="score-value">{recommendation}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_profile_cards(items: List[tuple[str, Any]], columns_per_row: int = 2) -> None:
    safe_columns = max(1, columns_per_row)
    for start in range(0, len(items), safe_columns):
        row = items[start : start + safe_columns]
        cols = st.columns(safe_columns, gap="small")
        for col, (label, value) in zip(cols, row):
            display = "—" if value in (None, "", "—") else value
            with col:
                st.markdown(
                    f"""
                    <div class="profile-card">
                        <div class="profile-label">{label}</div>
                        <div class="profile-value">{display}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )


def render_soft_card_grid(items: List[Dict[str, str]], columns_per_row: int = 2, top_margin: str = "0.85rem") -> None:
    safe_columns = max(1, columns_per_row)
    for start in range(0, len(items), safe_columns):
        row = items[start : start + safe_columns]
        cols = st.columns(safe_columns, gap="large")
        for col, item in zip(cols, row):
            with col:
                st.markdown(
                    f"""
                    <div class="soft-card" style="margin-top:{top_margin};">
                        <div class="mini-label">{item['title']}</div>
                        <div class="board-name" style="font-size:1rem;">{item['winner']}</div>
                        <div class="memo-text" style="margin-top:0.45rem;">{item['note']}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )


def render_five_layer_diagnostic(detail: Dict[str, Any]) -> None:
    st.markdown('<div class="section-kicker" style="margin-top:1rem;">Diagnostic Layer</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Five Layer Diagnostic</div>', unsafe_allow_html=True)
    for row in compute_five_layer_diagnostic(detail):
        st.markdown(
            f"""
            <div class="soft-card" style="margin-bottom:0.75rem;">
                <div class="diagnostic-row-head">
                    <div class="diagnostic-row-grade">{row['grade']}</div>
                    <div class="mini-label" style="margin-bottom:0;">{row['layer']}</div>
                </div>
                <div class="memo-text" style="margin-top:0.45rem;">{row['note']}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_front_office_workspace(detail: Dict[str, Any]) -> None:
    evaluation_id = str(detail.get("id") or "")
    if not evaluation_id:
        return

    workflow = get_front_office_meta(detail)
    st.markdown('<div class="section-kicker" style="margin-top:0.9rem;">Front-Office Workflow</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Shadow Board Controls</div>', unsafe_allow_html=True)

    summary_cards = [
        {"title": "Board Stage", "winner": workflow.get("stage", "Shadow Board"), "note": workflow.get("next_action", "No next step set.")},
        {"title": "Value Posture", "winner": workflow.get("value_posture", "Price-sensitive"), "note": workflow.get("spend_discipline", "Keep the acquisition disciplined.")},
        {"title": "Market Band", "winner": workflow.get("market_band", "Market"), "note": f"{workflow.get('priority_tier', 'Tier 3')} file for the current board."},
    ]
    render_soft_card_grid(summary_cards, columns_per_row=3, top_margin="0")

    control_left, control_mid, control_right = st.columns([0.75, 0.65, 1.3], gap="small")
    with control_left:
        stage = st.selectbox(
            "Board stage",
            WORKFLOW_STAGES,
            index=max(0, WORKFLOW_STAGES.index(workflow.get("stage", "Shadow Board"))),
            key=f"workflow_stage_{evaluation_id}",
        )
    with control_mid:
        owner = st.text_input(
            "Decision owner",
            value=str(workflow.get("owner") or "GM"),
            key=f"workflow_owner_{evaluation_id}",
        )
    with control_right:
        next_action = st.text_input(
            "Next action",
            value=str(workflow.get("next_action") or ""),
            key=f"workflow_next_{evaluation_id}",
        )

    save_front_office_meta(
        evaluation_id,
        stage=stage,
        owner=owner.strip() or workflow.get("owner") or "GM",
        next_action=next_action.strip() or workflow.get("next_action") or "",
    )
    st.caption("This is the lightweight operating layer for smaller staffs: stage the file, assign an owner, and define the next move.")


def build_board_filter_summary(
    preferred_mode: str,
    mode_scope: str,
    stage_filter: str,
    action_filter: str,
) -> str:
    parts: List[str] = []
    if mode_scope == "Focused":
        parts.append(f"Mode: {MODE_LABELS.get(preferred_mode, preferred_mode)}")
    elif mode_scope != "All":
        parts.append(f"Mode: {MODE_LABELS.get(mode_scope, mode_scope)}")
    else:
        parts.append("Mode: All")

    parts.append(f"Stage: {stage_filter}")
    parts.append(f"Recommendation: {action_filter}")
    return " | ".join(parts)


def build_budget_scenarios(detail: Dict[str, Any]) -> List[Dict[str, str]]:
    detail = normalize_detail_for_display(detail)
    player = detail.get("player", {}) or {}
    workflow = get_front_office_meta(detail)
    score = _safe_float(detail.get("overall_score"))
    value_score = _safe_float((detail.get("components") or {}).get("value"), score)
    cost_tier = _safe_float(player.get("expected_cost_tier"))
    mode = detail.get("mode") or DEFAULT_MODE

    walkaway_threshold = {
        "pro_wnba": max(2.0, cost_tier + 0.5),
        "cbb_high_major": max(2.0, cost_tier + 0.75),
        "cbb_d2_low_resource": max(1.0, cost_tier + 0.25),
        "recruiting_only": max(1.5, cost_tier + 0.5),
    }.get(mode, cost_tier + 0.5)

    if value_score >= 78 and cost_tier <= 2:
        disciplined_note = "This is the kind of file a lower-resource staff can actually land without breaking the board."
    elif value_score >= 72:
        disciplined_note = "Playable price if role certainty stays high. Do not turn this into a market-chasing fight."
    else:
        disciplined_note = "Only works if the market softens or the roster need becomes urgent."

    stretch_note = (
        "Stretch only if the staff sees a true top-7 rotation answer, not just a playable add."
        if score >= 76
        else "Stretching here is mostly paying for scarcity or panic rather than clean value."
    )
    walkaway_note = (
        f"Walk away once the file moves above cost tier {walkaway_threshold:.1f}. "
        f"{workflow.get('value_posture', 'Price-sensitive')} should still drive the call."
    )

    return [
        {"title": "Disciplined Price", "winner": workflow.get("market_band", "Market"), "note": disciplined_note},
        {"title": "Stretch Case", "winner": "Stretch only with conviction", "note": stretch_note},
        {"title": "Walk-Away Line", "winner": f"Tier {walkaway_threshold:.1f}", "note": walkaway_note},
    ]


def render_budget_scenario_view(detail: Dict[str, Any]) -> None:
    st.markdown('<div class="section-kicker" style="margin-top:0.9rem;">Budget Discipline</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Budget Scenario View</div>', unsafe_allow_html=True)
    render_soft_card_grid(build_budget_scenarios(detail), columns_per_row=3, top_margin="0")


def render_medical_diligence_workspace(detail: Dict[str, Any], editable: bool = True) -> None:
    evaluation_id = str(detail.get("id") or "")
    if not evaluation_id:
        return

    diligence = get_front_office_meta(detail)
    role = current_workspace_role()
    st.markdown(f'<div class="section-kicker" style="margin-top:0.9rem;">Shared Staff Overlay</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="section-title">{MEDICAL_DILIGENCE_LABEL}</div>', unsafe_allow_html=True)
    if editable:
        render_workspace_notice(MEDICAL_DILIGENCE_LABEL, evaluation_id)

    cards = [
        {"title": "Diligence Call", "winner": diligence.get("level", "Proceed with caution"), "note": diligence.get("gm_note", "No note entered.") or "No GM-facing note entered yet."},
        {"title": "Confidence", "winner": diligence.get("confidence", "Low"), "note": diligence.get("source", "Awaiting public-file review")},
        {"title": "Movement / Availability", "winner": "Review needed", "note": diligence.get("movement_flag", "No movement review entered.") or "No movement review entered yet."},
        {
            "title": "Report Status",
            "winner": "Complete" if diligence.get("reviewed_at") else "Open",
            "note": (
                f"Last saved {format_dt(diligence.get('reviewed_at'))} by {diligence.get('reviewed_by') or 'Staff'}"
                if diligence.get("reviewed_at")
                else "No Med Diligence report saved yet."
            ),
        },
    ]
    render_soft_card_grid(cards, columns_per_row=4, top_margin="0")

    history = get_front_office_audit(evaluation_id)
    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0.8rem;">
            <div class="mini-label">Public Injury / Missed-Time History</div>
            <div class="memo-text">{diligence.get('public_history') or 'No public history summary entered yet.'}</div>
            <div class="mini-label" style="margin-top:0.8rem;">Reviewed By</div>
            <div class="memo-text">{diligence.get('reviewed_by') or 'No staff review logged yet.'}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if not editable:
        with st.expander("Med Diligence History", expanded=False):
            if not history:
                st.caption("No Med Diligence saves logged yet.")
            else:
                for entry in reversed(history):
                    st.markdown(
                        f"**{entry['actor']}** ? {format_dt(entry['at'])}<br><span style='color:#64748b'>{entry['summary']}</span>",
                        unsafe_allow_html=True,
                    )
        return

    with st.expander(f"Edit {MEDICAL_DILIGENCE_LABEL} Report", expanded=False):
        render_workspace_notice(MEDICAL_DILIGENCE_LABEL, evaluation_id)
        c1, c2 = st.columns([0.8, 0.7], gap="small")
        with c1:
            level = st.selectbox(
                f"{MEDICAL_DILIGENCE_LABEL} call",
                MEDICAL_DILIGENCE_LEVELS,
                index=max(0, MEDICAL_DILIGENCE_LEVELS.index(diligence.get("level", MEDICAL_DILIGENCE_LEVELS[1]))),
                key=f"med_level_{evaluation_id}",
            )
        with c2:
            confidence = st.selectbox(
                "Confidence",
                MEDICAL_CONFIDENCE_LEVELS,
                index=max(0, MEDICAL_CONFIDENCE_LEVELS.index(diligence.get("confidence", "Low"))),
                key=f"med_conf_{evaluation_id}",
            )

        movement_flag = st.text_input(
            "Movement / functional flag",
            value=str(diligence.get("movement_flag") or ""),
            key=f"med_move_{evaluation_id}",
        )
        public_history = st.text_input(
            "Public injury / missed-time history",
            value=str(diligence.get("public_history") or ""),
            key=f"med_hist_{evaluation_id}",
        )
        gm_note = st.text_input(
            "GM-facing diligence note",
            value=str(diligence.get("gm_note") or ""),
            key=f"med_note_{evaluation_id}",
        )
        c3, c4 = st.columns([0.7, 1.3], gap="small")
        with c3:
            reviewed_by = st.text_input(
                "Reviewed by",
                value=str(diligence.get("reviewed_by") or WORKSPACE_ROLES.get(role, "Staff")),
                key=f"med_reviewed_by_{evaluation_id}",
            )
        with c4:
            st.markdown(
                """
                <div class="soft-card" style="margin-top:0;">
                    <div class="mini-label">Privacy Guardrail</div>
                    <div class="memo-text">Use only public history, public video, and staff movement observations here. Do not enter protected student medical records or treatment details.</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        if st.button("Save Med Diligence Report", key=f"save_med_diligence_{evaluation_id}", width="stretch"):
            actor = reviewed_by.strip() or WORKSPACE_ROLES.get(role, "Staff")
            clean_level = level
            clean_confidence = confidence
            clean_movement = movement_flag.strip()
            clean_history = public_history.strip()
            clean_note = gm_note.strip()
            save_front_office_meta(
                evaluation_id,
                level=clean_level,
                confidence=clean_confidence,
                movement_flag=clean_movement,
                public_history=clean_history,
                gm_note=clean_note,
                reviewed_by=actor,
                reviewed_at=_now_iso(),
                source=f"{WORKSPACE_ROLES.get(role, 'Staff')} public-file review",
            )
            summary = f"{clean_level} | {clean_confidence}"
            if clean_note:
                summary += f" | {clean_note}"
            append_front_office_audit(evaluation_id, MEDICAL_DILIGENCE_LABEL, actor, summary)
            push_workspace_notice(MEDICAL_DILIGENCE_LABEL, evaluation_id, "Med Diligence report saved. Status is now complete for this file.")
            st.rerun()
        st.caption(
            "Advisory diligence only. Keep this limited to public-file review, movement observations, and roster-risk framing."
        )
    with st.expander("Med Diligence History", expanded=False):
        if not history:
            st.caption("No Med Diligence saves logged yet.")
        else:
            for entry in reversed(history):
                st.markdown(
                    f"**{entry['actor']}** ? {format_dt(entry['at'])}<br><span style='color:#64748b'>{entry['summary']}</span>",
                    unsafe_allow_html=True,
                )


def _evidence_lines(raw_text: str) -> List[str]:
    return [line.strip() for line in (raw_text or "").splitlines() if line.strip()]


def _parse_evidence_line(line: str) -> tuple[str, str, str]:
    parts = [part.strip() for part in line.split("|", 2)]
    if len(parts) == 3:
        return parts[0], parts[1], parts[2]
    if len(parts) == 2:
        return parts[0], parts[1], ""
    return line.strip(), "", ""


def build_verified_source_snapshot(detail: Dict[str, Any]) -> Dict[str, str]:
    meta = get_front_office_meta(detail)
    evidence_lines = _evidence_lines(str(meta.get("research_evidence_log") or meta.get("research_source_links") or ""))
    if not evidence_lines:
        return {
            "title": "Verified Source Snapshot",
            "winner": "Not logged",
            "note": "No official NCAA, ESPN, school, or public evidence source has been attached yet.",
        }

    label, url, note = _parse_evidence_line(evidence_lines[0])
    winner = label or "Source logged"
    note_parts = []
    if url:
        note_parts.append(url)
    if note:
        note_parts.append(note)
    if len(evidence_lines) > 1:
        note_parts.append(f"{len(evidence_lines)} total sources logged")
    return {
        "title": "Verified Source Snapshot",
        "winner": winner,
        "note": " · ".join(note_parts) if note_parts else "Public evidence source logged for this file.",
    }


def render_research_evidence_log(detail: Dict[str, Any]) -> None:
    meta = get_front_office_meta(detail)
    evidence_lines = _evidence_lines(str(meta.get("research_evidence_log") or meta.get("research_source_links") or ""))
    st.markdown('<div class="section-kicker" style="margin-top:0.4rem;">Source Log</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Research Evidence</div>', unsafe_allow_html=True)
    if not evidence_lines:
        st.info("No evidence sources logged yet.")
        return
    for line in evidence_lines:
        parts = [part.strip() for part in line.split("|", 2)]
        if len(parts) == 3:
            label, url, note = parts
            body = f"<strong>{label}</strong><br><span style='color:#64748b'>{url}</span><br>{note}"
        elif len(parts) == 2:
            label, url = parts
            body = f"<strong>{label}</strong><br><span style='color:#64748b'>{url}</span>"
        else:
            body = line
        st.markdown(f"<div class='soft-card' style='margin-top:0.5rem;'><div class='memo-text'>{body}</div></div>", unsafe_allow_html=True)


def render_completed_reports(detail: Dict[str, Any]) -> None:
    meta = get_front_office_meta(detail)
    evaluation_id = str(detail.get("id") or "")
    history = get_front_office_audit(evaluation_id)
    research_complete = bool(meta.get("research_summary"))
    med_complete = bool(meta.get("reviewed_at"))
    last_research = get_latest_audit_entry(evaluation_id, "Prospect Research")
    last_med = get_latest_audit_entry(evaluation_id, MEDICAL_DILIGENCE_LABEL)
    st.markdown('<div class="section-kicker" style="margin-top:0.4rem;">Completed Staff Work</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Completed Reports</div>', unsafe_allow_html=True)
    render_soft_card_grid(
        [
            {
                "title": "Prospect Research",
                "winner": "Complete" if research_complete else "Open",
                "note": (
                    f"Last saved {format_dt(last_research.get('at'))} by {last_research.get('actor')}"
                    if last_research
                    else "No research report saved yet."
                ),
            },
            {
                "title": MEDICAL_DILIGENCE_LABEL,
                "winner": "Complete" if med_complete else "Open",
                "note": (
                    f"Last saved {format_dt(last_med.get('at'))} by {last_med.get('actor')}"
                    if last_med
                    else "No Med Diligence report saved yet."
                ),
            },
        ],
        columns_per_row=2,
        top_margin="0",
    )
    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0;">
            <div class="mini-label">Research Summary</div>
            <div class="memo-text">{meta.get('research_summary') or 'No completed research summary saved yet.'}</div>
            <div class="mini-label" style="margin-top:0.8rem;">Recommendation To GM</div>
            <div class="memo-text">{meta.get('research_recommendation') or meta.get('gm_note') or 'No recommendation saved yet.'}</div>
            <div class="mini-label" style="margin-top:0.8rem;">Med Diligence Status</div>
            <div class="memo-text">{meta.get('level') or 'No Med Diligence report saved yet.'}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.expander("Audit History", expanded=False):
        if not history:
            st.caption("No saved staff reports logged yet.")
        else:
            for entry in reversed(history):
                st.markdown(
                    f"**{entry['section']}** ? **{entry['actor']}** ? {format_dt(entry['at'])}<br><span style='color:#64748b'>{entry['summary']}</span>",
                    unsafe_allow_html=True,
                )


def build_staff_reports_summary(detail: Dict[str, Any]) -> List[Dict[str, str]]:
    meta = get_front_office_meta(detail)
    evaluation_id = str(detail.get("id") or "")
    last_research = get_latest_audit_entry(evaluation_id, "Prospect Research")
    last_med = get_latest_audit_entry(evaluation_id, MEDICAL_DILIGENCE_LABEL)
    evidence_count = len(_evidence_lines(str(meta.get("research_evidence_log") or meta.get("research_source_links") or "")))
    return [
        {
            "title": "Prospect Research",
            "winner": "Complete" if meta.get("research_summary") else "Open",
            "note": (
                f"Last saved {format_dt(last_research.get('at'))} by {last_research.get('actor')}"
                if last_research
                else "No saved research report yet."
            ),
        },
        {
            "title": MEDICAL_DILIGENCE_LABEL,
            "winner": "Complete" if meta.get("reviewed_at") else "Open",
            "note": (
                f"Last saved {format_dt(last_med.get('at'))} by {last_med.get('actor')}"
                if last_med
                else "No saved Med Diligence report yet."
            ),
        },
        {
            "title": "Evidence Sources",
            "winner": str(evidence_count),
            "note": "Public-source links or notes attached to this file.",
        },
    ]


def render_prospect_research_workspace(detail: Dict[str, Any], editable: bool = True) -> None:
    evaluation_id = str(detail.get("id") or "")
    if not evaluation_id:
        return

    meta = get_front_office_meta(detail)
    role = current_workspace_role()
    source_links = str(meta.get("research_source_links") or "")
    evidence_log = str(meta.get("research_evidence_log") or source_links)
    video_links = str(meta.get("research_video_links") or "")
    injury_history = str(meta.get("research_injury_history") or meta.get("public_history") or "")
    movement_notes = str(meta.get("research_movement_notes") or meta.get("movement_flag") or "")
    research_summary = str(meta.get("research_summary") or "")
    recommendation = str(meta.get("research_recommendation") or "")
    owner = str(meta.get("research_owner") or WORKSPACE_ROLES.get(role, "Staff"))
    research_audit = get_latest_audit_entry(evaluation_id, "Prospect Research")

    st.markdown('<div class="section-kicker" style="margin-top:0.9rem;">External Prospect Review</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Prospect Research Workspace</div>', unsafe_allow_html=True)
    st.caption(
        "Use this for outside-player diligence: public history, available video, movement observations, and a recommendation back to the GM file."
    )
    render_soft_card_grid(
        [
            {
                "title": "Report Status",
                "winner": "Complete" if research_summary else "Open",
                "note": (
                    f"Last saved {format_dt(research_audit.get('at'))} by {research_audit.get('actor')}"
                    if research_audit
                    else "No prospect research report saved yet."
                ),
            },
            {
                "title": "Research Owner",
                "winner": owner or "Staff",
                "note": recommendation or "No GM recommendation saved yet.",
            },
        ],
        columns_per_row=2,
        top_margin="0",
    )

    if not editable:
        st.markdown(
            f"""
            <div class="soft-card" style="margin-top:0;">
                <div class="mini-label">Research Owner</div>
                <div class="memo-text">{owner}</div>
                <div class="mini-label" style="margin-top:0.8rem;">Research Summary</div>
                <div class="memo-text">{research_summary or 'No research summary saved yet.'}</div>
                <div class="mini-label" style="margin-top:0.8rem;">Recommendation To GM</div>
                <div class="memo-text">{recommendation or 'No recommendation saved yet.'}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        return

    with st.expander("Edit Prospect Research Report", expanded=False):
        render_workspace_notice("Prospect Research", evaluation_id)
        c1, c2 = st.columns(2, gap="small")
        with c1:
            research_owner = st.text_input(
                "Review owner",
                value=owner,
                key=f"research_owner_{evaluation_id}",
            )
            evidence_log_text = st.text_area(
                "Research evidence log",
                value=evidence_log,
                height=120,
                key=f"research_evidence_{evaluation_id}",
                placeholder="Source label | URL | What this source supports",
            )
            public_history_text = st.text_area(
                "Public injury / missed-time history",
                value=injury_history,
                height=110,
                key=f"research_history_{evaluation_id}",
                placeholder="Summarize only what is public or staff-observed. Do not paste protected medical records.",
            )
        with c2:
            video_review_links = st.text_area(
                "Video / movement review links",
                value=video_links,
                height=100,
                key=f"research_videos_{evaluation_id}",
                placeholder="Synergy clips, broadcast links, YouTube, school film...",
            )
            movement_review = st.text_area(
                "Movement / functional observations",
                value=movement_notes,
                height=110,
                key=f"research_movement_{evaluation_id}",
                placeholder="How the player moves, changes direction, lands, accelerates, or manages workload on film.",
            )

        research_summary_text = st.text_area(
            "Research summary",
            value=research_summary,
            height=100,
            key=f"research_summary_{evaluation_id}",
            placeholder="Concise synthesis of what the staff believes after public-file and video review.",
        )
        research_recommendation = st.text_input(
            "Recommendation to GM",
            value=recommendation,
            key=f"research_reco_{evaluation_id}",
            placeholder="Example: Proceed, but keep pricing disciplined until deeper diligence is complete.",
        )

        if st.button("Save Prospect Research Report", key=f"save_research_{evaluation_id}", width="stretch"):
            actor = research_owner.strip() or WORKSPACE_ROLES.get(role, "Staff")
            save_front_office_meta(
                evaluation_id,
                research_owner=actor,
                research_source_links=evidence_log_text.strip(),
                research_evidence_log=evidence_log_text.strip(),
                research_video_links=video_review_links.strip(),
                research_injury_history=public_history_text.strip(),
                research_movement_notes=movement_review.strip(),
                research_summary=research_summary_text.strip(),
                research_recommendation=research_recommendation.strip(),
                public_history=public_history_text.strip() or meta.get("public_history") or "",
                movement_flag=movement_review.strip() or meta.get("movement_flag") or "",
                gm_note=research_recommendation.strip() or meta.get("gm_note") or "",
                source=f"{WORKSPACE_ROLES.get(role, 'Staff')} public-file review",
            )
            summary = research_recommendation.strip() or research_summary_text.strip() or "Research updated"
            append_front_office_audit(evaluation_id, "Prospect Research", actor, summary)
            push_workspace_notice("Prospect Research", evaluation_id, "Prospect research report saved. Status is now complete for this file.")
            st.rerun()
        st.caption("Keep sources public, sport-facing, and advisory. This workspace should inform pricing and board stage, not replace formal medical clearance.")


def render_prospect_request_workspace(detail: Dict[str, Any]) -> None:
    evaluation_id = str(detail.get("id") or "")
    if not evaluation_id:
        return

    workflow = get_front_office_meta(detail)
    request_defaults = build_prospect_diligence_request(detail)

    st.markdown('<div class="section-kicker" style="margin-top:0.9rem;">Cross-Product Workflow</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Prospect Diligence Request</div>', unsafe_allow_html=True)
    st.caption("Use this when the player is outside WAIMS Python and you need sport science / medical staff to do a public-file diligence review for the GM board.")

    left, right = st.columns(2, gap="small")
    with left:
        requested_by = st.text_input(
            "Requested by",
            value=str(request_defaults.get("requested_by") or workflow.get("owner") or "GM"),
            key=f"prospect_request_by_{evaluation_id}",
        )
        request_focus = st.text_input(
            "Request focus",
            value=str(request_defaults.get("request_focus") or "External prospect diligence"),
            key=f"prospect_request_focus_{evaluation_id}",
        )
    with right:
        priority_tier = st.text_input(
            "Priority tier",
            value=str(request_defaults.get("priority_tier") or workflow.get("priority_tier") or "Tier 3"),
            key=f"prospect_request_priority_{evaluation_id}",
        )
    questions = st.text_area(
        "Questions for staff",
        value=str(request_defaults.get("questions") or ""),
        height=90,
        key=f"prospect_request_questions_{evaluation_id}",
    )
    gm_context = st.text_area(
        "GM context",
        value=str(request_defaults.get("gm_context") or ""),
        height=90,
        key=f"prospect_request_context_{evaluation_id}",
    )

    request_payload = build_prospect_diligence_request(
        detail,
        overrides={
            "requested_by": requested_by.strip(),
            "request_focus": request_focus.strip(),
            "questions": questions.strip(),
            "gm_context": gm_context.strip(),
        },
    )
    request_json = json.dumps(request_payload, indent=2)

    action_left, action_right = st.columns(2, gap="small")
    with action_left:
        st.download_button(
            "Download request (.json)",
            data=request_json,
            file_name=f"{((detail.get('player') or {}).get('name', 'player')).replace(' ', '_').lower()}_prospect_request.json",
            mime="application/json",
            use_container_width=True,
        )
    with action_right:
        if st.button("Send to WAIMS Python queue", key=f"send_request_{evaluation_id}", use_container_width=True):
            PROSPECT_REQUEST_DIR.mkdir(parents=True, exist_ok=True)
            target = PROSPECT_REQUEST_DIR / f"{request_payload['player_id'].lower()}_{request_payload['team_id'].lower()}.json"
            target.write_text(request_json, encoding="utf-8")
            st.success(f"Queued staff diligence request: {target}")


def render_detail(detail: Dict[str, Any], show_diagnostic: bool = True, show_heading: bool = True) -> None:
    detail = normalize_detail_for_display(detail)
    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    components = detail.get("components", {}) or {}
    assumptions = detail.get("assumptions", {}) or {}
    tension_points = detail.get("tension_points", []) or []
    mode = detail.get("mode") or "pro_wnba"

    if show_heading:
        st.markdown('<div class="section-kicker">Selected File</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-title">Evaluation Dossier</div>', unsafe_allow_html=True)

    st.markdown(
        f"""
        <div class="detail-shell">
            <div class="board-head">
                <div>
                    <div class="player-title">{player.get("name", "Unknown Player")}</div>
                    <div class="player-meta">
                        {MODE_LABELS.get(mode, mode)} &nbsp;|&nbsp;
                        {player.get("position", "—")} &nbsp;|&nbsp;
                        Team {detail.get("team_id", "—")} &nbsp;|&nbsp;
                        Score {format_score(detail.get("overall_score"))} &nbsp;|&nbsp;
                        {format_dt(detail.get("created_at"))}
                    </div>
                </div>
                <div class="board-tag {action_class(detail.get("recommended_action"))}">{clean_action(detail.get("recommended_action"), mode)}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    render_diagnostic_strip(detail)
    st.markdown('<div class="rule"></div>', unsafe_allow_html=True)

    render_soft_card_grid(build_executive_brief_cards(detail), columns_per_row=3, top_margin="0")
    render_level_delta_section(detail)
    takeaway_items = "".join(f"<li>{item}</li>" for item in build_dossier_takeaways(detail))
    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0.8rem; margin-bottom:0.9rem;">
            <div class="mini-label">Top Takeaways</div>
            <ul class="subtle-list">{takeaway_items}</ul>
        </div>
        """,
        unsafe_allow_html=True,
    )
    render_soft_card_grid([build_archetype_profile_card(detail)], columns_per_row=1, top_margin="0")
    with st.expander("Full Evaluation Breakdown", expanded=False):
        top_left, top_right = st.columns([0.95, 1.2], gap="large")

        with top_left:
            render_profile_cards(
                [
                    ("Age", player.get("age", "—")),
                    ("Position", player.get("position", "—")),
                    ("Cost Tier", player.get("expected_cost_tier", "—")),
                    ("Health Risk", player.get("health_risk", "—")),
                    ("Upside", player.get("upside", "—")),
                    ("Minutes Stability", player.get("minutes_stability", "—")),
                ]
            )

            st.markdown(
                f"""
                <div class="soft-card" style="margin-top:0.85rem;">
                    <div class="mini-label">Context Summary</div>
                    <div class="memo-text">{summarize_context(ctx)}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        with top_right:
            st.markdown(
                f"""
                <div class="soft-card">
                    <div class="mini-label">Executive Memo</div>
                    <div class="memo-text">{build_memo(detail)}</div>
                    <div class="rule"></div>
                    <div class="mini-label">Decision Lens</div>
                    <div class="memo-text">{build_decision_lens(detail)}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        mid_left, mid_right = st.columns([0.95, 1.2], gap="large")

        with mid_left:
            render_score_cards(detail, components)

            assumption_items = "".join(
                [f"<li>{k.replace('_', ' ').title()}: {v}</li>" for k, v in assumptions.items()]
            ) or "<li>No assumptions recorded.</li>"

            st.markdown(
                f"""
                <div class="soft-card" style="margin-top:0.85rem;">
                    <div class="mini-label">Assumptions</div>
                    <ul class="subtle-list">{assumption_items}</ul>
                </div>
                """,
                unsafe_allow_html=True,
            )

        with mid_right:
            needs = ctx.get("needs_by_position", {}) or {}
            needs_text = ", ".join(f"{k}: {v}" for k, v in needs.items()) if needs else "No positional needs recorded."
            tension_items = "".join([f"<li>{item}</li>" for item in tension_points]) or "<li>No major tension points flagged.</li>"

            st.markdown(
                f"""
                <div class="soft-card">
                    <div class="mini-label">Positional Needs</div>
                    <div class="memo-text">{needs_text}</div>
                    <div class="rule"></div>
                    <div class="mini-label">Tension Points</div>
                    <ul class="subtle-list">{tension_items}</ul>
                </div>
                """,
                unsafe_allow_html=True,
            )

        bottom_left, bottom_right = st.columns([1, 1], gap="large")

        with bottom_left:
            st.markdown(
                f"""
                <div class="soft-card">
                    <div class="mini-label">Strengths</div>
                    <ul class="subtle-list">{render_bullets(detail.get("strengths"), "No strengths entered.")}</ul>
                </div>
                """,
                unsafe_allow_html=True,
            )

        with bottom_right:
            st.markdown(
                f"""
                <div class="soft-card">
                    <div class="mini-label">Concerns</div>
                    <ul class="subtle-list">{render_bullets(detail.get("concerns"), "No concerns entered.")}</ul>
                </div>
                """,
                unsafe_allow_html=True,
            )

    if show_diagnostic:
        with st.expander("Five Layer Diagnostic", expanded=False):
            render_five_layer_diagnostic(detail)


def render_compare_block(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> None:
    left_detail = normalize_detail_for_display(left_detail)
    right_detail = normalize_detail_for_display(right_detail)
    left_lens = derive_level_delta(left_detail)
    right_lens = derive_level_delta(right_detail)

    st.markdown('<div class="section-kicker" style="margin-top:1rem;">Comparison Layer</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Side-by-Side Comparison</div>', unsafe_allow_html=True)

    left_name = (left_detail.get("player") or {}).get("name", "Selected Player")
    right_name = (right_detail.get("player") or {}).get("name", "Comparison Player")
    left_components = left_detail.get("components", {}) or {}
    right_components = right_detail.get("components", {}) or {}

    c1, c2 = st.columns(2, gap="large")
    with c1:
        st.markdown(f"**{left_name}**")
        st.metric("Overall Score", format_score(left_detail.get("overall_score")))
        st.metric("Action", clean_action(left_detail.get("recommended_action"), left_detail.get("mode")))
        st.write(left_detail.get("summary_note") or "No summary note.")
    with c2:
        st.markdown(f"**{right_name}**")
        st.metric("Overall Score", format_score(right_detail.get("overall_score")))
        st.metric("Action", clean_action(right_detail.get("recommended_action"), right_detail.get("mode")))
        st.write(right_detail.get("summary_note") or "No summary note.")

    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0.85rem;">
            <div class="mini-label">Comparison Summary</div>
            <div class="memo-text">{compare_summary(left_detail, right_detail)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    render_soft_card_grid(
        [
            {
                "title": f"{left_name} Bet Type",
                "winner": left_lens["bet_label"],
                "note": f"{left_lens['level_label']} | {left_lens['delta_label']}",
            },
            {
                "title": f"{right_name} Bet Type",
                "winner": right_lens["bet_label"],
                "note": f"{right_lens['level_label']} | {right_lens['delta_label']}",
            },
        ],
        columns_per_row=2,
        top_margin="0",
    )

    roster_need_call = build_roster_need_call(left_detail, right_detail)
    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0.85rem;">
            <div class="mini-label">{roster_need_call['title']}</div>
            <div class="board-name" style="font-size:1.05rem;">{roster_need_call['winner']}</div>
            <div class="memo-text" style="margin-top:0.45rem;">{roster_need_call['note']}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    render_soft_card_grid([build_archetype_comparison_card(left_detail, right_detail)], columns_per_row=1)

    decision_snapshot = build_compare_decision_snapshot(left_detail, right_detail)
    if decision_snapshot:
        render_soft_card_grid(decision_snapshot, columns_per_row=2)

    verdicts = build_comparison_verdicts(left_detail, right_detail)
    if verdicts:
        render_soft_card_grid(verdicts, columns_per_row=2)

    comparison_rows = []
    for key, label in COMPONENT_LABELS.items():
        left_value = left_components.get(key)
        right_value = right_components.get(key)

        if left_value is None and right_value is None:
            continue

        if left_value is None:
            edge = right_name
        elif right_value is None:
            edge = left_name
        elif abs(float(left_value) - float(right_value)) < 0.01:
            edge = "Even"
        else:
            edge = left_name if float(left_value) > float(right_value) else right_name

        comparison_rows.append(
            f"""
            <div class="compare-row">
                <div class="compare-metric">{label}</div>
                <div class="compare-score">{format_score(left_value)}</div>
                <div class="compare-advantage">{edge}</div>
                <div class="compare-score">{format_score(right_value)}</div>
            </div>
            """
        )

    if comparison_rows:
        st.markdown('<div class="section-kicker" style="margin-top:1rem;">Component Layer</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-title">Component Comparison</div>', unsafe_allow_html=True)
        st.markdown(
            f"""
            <div class="compare-grid">
                {''.join(comparison_rows)}
            </div>
            """,
            unsafe_allow_html=True,
        )

    notes_left, notes_right = st.columns(2, gap="large")
    with notes_left:
        st.markdown(
            f"""
            <div class="soft-card" style="margin-top:0.85rem;">
                <div class="mini-label">{left_name} Strengths</div>
                <ul class="subtle-list">{render_bullets(left_detail.get("strengths"), "No strengths entered.")}</ul>
                <div class="rule"></div>
                <div class="mini-label">{left_name} Concerns</div>
                <ul class="subtle-list">{render_bullets(left_detail.get("concerns"), "No concerns entered.")}</ul>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with notes_right:
        st.markdown(
            f"""
            <div class="soft-card" style="margin-top:0.85rem;">
                <div class="mini-label">{right_name} Strengths</div>
                <ul class="subtle-list">{render_bullets(right_detail.get("strengths"), "No strengths entered.")}</ul>
                <div class="rule"></div>
                <div class="mini-label">{right_name} Concerns</div>
                <ul class="subtle-list">{render_bullets(right_detail.get("concerns"), "No concerns entered.")}</ul>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_collaborator_home(evaluations: List[Dict[str, Any]], mode_scope: str, stage_filter: str, action_filter: str) -> None:
    role = current_workspace_role()
    role_label = WORKSPACE_ROLES.get(role, role)
    st.markdown('<div class="section-kicker">Collaborator Workspace</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Staff Landing Page</div>', unsafe_allow_html=True)
    st.caption("This workspace is for public-file prospect research and Med Diligence only. Board control, pricing, and create/delete actions stay with the GM.")

    active_items = len(evaluations)
    ready_for_review = sum(1 for row in evaluations if (get_front_office_meta(row).get("stage") or "Shadow Board") in {"Priority", "Live Board"})
    completed_staff = sum(1 for row in evaluations if get_front_office_meta(row).get("reviewed_at") or get_front_office_meta(row).get("research_summary"))
    render_soft_card_grid(
        [
            {"title": "Signed In As", "winner": role_label, "note": "Collaborator role"},
            {"title": "Visible Files", "winner": str(active_items), "note": "Current filtered workspace"},
            {"title": "Priority / Live", "winner": str(ready_for_review), "note": "Closest to active review"},
            {"title": "Completed Staff Notes", "winner": str(completed_staff), "note": "Files with saved research or diligence"},
        ],
        columns_per_row=4,
        top_margin="0",
    )

    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0.8rem;">
            <div class="mini-label">How To Use This View</div>
            <div class="memo-text">
                1. Use the sidebar filters to narrow the board.<br>
                2. Open <strong>Player Dossier</strong> and use the dossier tabs for <strong>Prospect Research</strong> and <strong>{MEDICAL_DILIGENCE_LABEL}</strong> authoring.<br>
                3. Keep all notes limited to public sources, staff film review, and roster-risk framing.<br>
                4. Completed writeups roll into the GM-facing readout automatically.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        f"""
        <div class="soft-card" style="margin-top:0.8rem;">
            <div class="mini-label">Active Board View</div>
            <div class="memo-text">{build_board_filter_summary(st.session_state.get("preferred_mode", DEFAULT_MODE), mode_scope, stage_filter, action_filter)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if not evaluations:
        st.info("No files match the current filters. Broaden the board filters in the sidebar or ask the GM to move more files into review.")
        return

    queue_rows = []
    for row in evaluations[:8]:
        meta = get_front_office_meta(row)
        queue_rows.append(
            {
                "Player": ((row.get("player") or {}).get("name") or "Unknown"),
                "Stage": meta.get("stage") or "Shadow Board",
                "Owner": meta.get("owner") or "GM",
                "Research": "Done" if meta.get("research_summary") else "Open",
                MEDICAL_DILIGENCE_LABEL: "Done" if meta.get("reviewed_at") else "Open",
            }
        )
    st.dataframe(queue_rows, hide_index=True, width="stretch")


def render_file_selector(
    evaluations: List[Dict[str, Any]],
    selected_id: Optional[str],
    token: str,
    preferred_mode: str,
    *,
    select_key: str,
    select_label: str = "Select player dossier",
    show_exports: bool = False,
) -> Optional[Dict[str, Any]]:
    if not evaluations:
        return None

    dossier_map = {
        f"{(e.get('player') or {}).get('name', 'Player')} | {MODE_LABELS.get(e.get('mode') or 'pro_wnba', e.get('mode') or 'pro_wnba')} | {format_score(e.get('overall_score'))}": e["id"]
        for e in evaluations
    }
    dossier_labels = list(dossier_map.keys())
    current_label = next((label for label, value in dossier_map.items() if value == selected_id), dossier_labels[0])

    if show_exports:
        control_left, control_mid, control_right = st.columns([1.8, 0.42, 0.42], gap="small")
    else:
        control_left, control_mid, control_right = st.columns([1.8, 0.42, 0.42], gap="small")

    with control_left:
        st.markdown('<div class="section-kicker" style="margin-bottom:0.2rem;">Select Player</div>', unsafe_allow_html=True)
        selected_label = st.selectbox(
            select_label,
            dossier_labels,
            index=dossier_labels.index(current_label),
            label_visibility="collapsed",
            key=select_key,
        )

    selected_id = dossier_map[selected_label]
    st.session_state["selected_evaluation_id"] = selected_id
    detail = get_evaluation_detail(token, selected_id)
    render_mode_focus_banner(detail.get("mode") or preferred_mode, "dossier", show_label=False)

    if show_exports:
        player_name = (detail.get("player") or {}).get("name", "player").replace(" ", "_")
        with control_mid:
            export_md = build_export_markdown(detail)
            st.download_button(
                "Markdown",
                data=export_md,
                file_name=f"{player_name}_waims_gm_dossier.md",
                mime="text/markdown",
                use_container_width=True,
                key=f"{select_key}_export_md",
            )
        with control_right:
            if WORD_EXPORT_AVAILABLE:
                try:
                    export_docx = build_export_docx_bytes(detail)
                    st.download_button(
                        "Word",
                        data=export_docx,
                        file_name=f"{player_name}_waims_gm_dossier.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"{select_key}_export_docx",
                    )
                except Exception as e:
                    st.warning(f"Word export failed, but the rest of the app is still available. Details: {e}")
            else:
                st.info("Word export is unavailable on this environment. Markdown export is still available.")

    return detail


def build_payload_from_form(
    display_name: str,
    player_id: str,
    player_name: str,
    position: str,
    age: int,
    offense_rating: float,
    defense_rating: float,
    shooting_rating: float,
    playmaking_rating: float,
    rebounding_rating: float,
    health_risk: float,
    upside: float,
    minutes_stability: float,
    expected_cost_tier: int,
    team_id: str,
    timeline: str,
    need_g: float,
    need_f: float,
    need_c: float,
    cap_flexibility: float,
    risk_tolerance: float,
    summary_note: str,
    strengths: str,
    concerns: str,
    mode: str,
) -> Dict[str, Any]:
    return {
        "player": {
            "id": player_id,
            "name": player_name,
            "position": position,
            "age": age,
            "offense_rating": offense_rating,
            "defense_rating": defense_rating,
            "shooting_rating": shooting_rating,
            "playmaking_rating": playmaking_rating,
            "rebounding_rating": rebounding_rating,
            "health_risk": health_risk,
            "upside": upside,
            "minutes_stability": minutes_stability,
            "expected_cost_tier": expected_cost_tier,
        },
        "ctx": {
            "team_id": team_id,
            "timeline": timeline,
            "needs_by_position": {"G": need_g, "F": need_f, "C": need_c},
            "cap_flexibility": cap_flexibility,
            "risk_tolerance": risk_tolerance,
        },
        "display_name": display_name,
        "summary_note": summary_note,
        "strengths": strengths,
        "concerns": concerns,
        "mode": mode,
    }


def render_edit_evaluation_workspace(token: str, detail: Dict[str, Any]) -> None:
    if current_workspace_role() != "gm":
        return

    detail = normalize_detail_for_display(detail)
    evaluation_id = str(detail.get("id") or "")
    if not evaluation_id:
        return

    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    mode = detail.get("mode") or DEFAULT_MODE
    team_presets = get_team_context_presets(mode)

    with st.expander("Edit Evaluation File", expanded=False):
        st.caption("Update the player file, team context, and scouting notes, then rescore the dossier without creating a duplicate.")
        edit_team_preset_name = st.selectbox(
            "Team Context Preset",
            ["Current"] + list(team_presets.keys()),
            index=0,
            key=f"edit_team_preset_{evaluation_id}",
        )
        edit_team_preset = {} if edit_team_preset_name == "Current" else team_presets.get(edit_team_preset_name, {})

        with st.form(f"edit_eval_form_{evaluation_id}"):
            st.markdown("**Player identity**")
            e1, e2 = st.columns(2)
            with e1:
                player_id = st.text_input("Player ID", value=str(player.get("id") or ""))
                player_name = st.text_input("Player name", value=str(player.get("name") or ""))
            with e2:
                position = st.selectbox("Position", ["G", "F", "C"], index=["G", "F", "C"].index(str(player.get("position") or "F")))
                age = st.number_input("Age", min_value=16, max_value=45, value=int(player.get("age") or 23), step=1)
                expected_cost_tier = st.number_input("Expected cost tier", min_value=0, max_value=10, value=int(player.get("expected_cost_tier") or 2), step=1)

            st.markdown("**Basketball profile**")
            e3, e4 = st.columns(2)
            with e3:
                offense_rating = st.slider("Offense", 0, 100, int(float(player.get("offense_rating") or 74)), key=f"edit_offense_{evaluation_id}")
                defense_rating = st.slider("Defense", 0, 100, int(float(player.get("defense_rating") or 74)), key=f"edit_defense_{evaluation_id}")
                shooting_rating = st.slider("Shooting", 0, 100, int(float(player.get("shooting_rating") or 74)), key=f"edit_shooting_{evaluation_id}")
            with e4:
                playmaking_rating = st.slider("Playmaking", 0, 100, int(float(player.get("playmaking_rating") or 60)), key=f"edit_playmaking_{evaluation_id}")
                rebounding_rating = st.slider("Rebounding", 0, 100, int(float(player.get("rebounding_rating") or 60)), key=f"edit_rebounding_{evaluation_id}")
                minutes_stability = st.slider("Minutes stability", 0.0, 1.0, float(player.get("minutes_stability") or 0.7), 0.01, key=f"edit_minutes_{evaluation_id}")

            st.markdown("**Risk and projection**")
            e5, e6 = st.columns(2)
            with e5:
                health_risk = st.slider("Health risk", 0.0, 1.0, float(player.get("health_risk") or 0.2), 0.01, key=f"edit_health_{evaluation_id}")
            with e6:
                upside = st.slider("Upside", 0.0, 1.0, float(player.get("upside") or 0.7), 0.01, key=f"edit_upside_{evaluation_id}")

            st.markdown("**Team context**")
            e7, e8 = st.columns(2)
            with e7:
                team_id = st.text_input("Team ID", value=str(edit_team_preset.get("team_id") or ctx.get("team_id") or "team-1"), key=f"edit_team_id_{evaluation_id}")
                timeline = st.selectbox(
                    "Timeline",
                    ["win_now", "balanced", "rebuild"],
                    index=["win_now", "balanced", "rebuild"].index(str(edit_team_preset.get("timeline") or ctx.get("timeline") or "balanced")),
                    key=f"edit_timeline_{evaluation_id}",
                )
                cap_flexibility = st.slider(
                    "Cap flexibility",
                    0.0,
                    1.0,
                    float(edit_team_preset.get("cap_flexibility", ctx.get("cap_flexibility", 0.60))),
                    0.01,
                    key=f"edit_cap_{evaluation_id}",
                )
            with e8:
                needs = ctx.get("needs_by_position", {}) or {}
                risk_tolerance = st.slider(
                    "Risk tolerance",
                    0.0,
                    1.0,
                    float(edit_team_preset.get("risk_tolerance", ctx.get("risk_tolerance", 0.40))),
                    0.01,
                    key=f"edit_risk_tol_{evaluation_id}",
                )
                need_g = st.slider("Need at Guard", 0.0, 1.0, float(edit_team_preset.get("need_g", needs.get("G", 0.55))), 0.01, key=f"edit_need_g_{evaluation_id}")
                need_f = st.slider("Need at Forward", 0.0, 1.0, float(edit_team_preset.get("need_f", needs.get("F", 0.80))), 0.01, key=f"edit_need_f_{evaluation_id}")
                need_c = st.slider("Need at Center", 0.0, 1.0, float(edit_team_preset.get("need_c", needs.get("C", 0.35))), 0.01, key=f"edit_need_c_{evaluation_id}")

            st.markdown("**Scouting rationale**")
            summary_note = st.text_area("Summary note", value=str(detail.get("summary_note") or ""), key=f"edit_summary_{evaluation_id}")
            strengths = st.text_area("Strengths", value=str(detail.get("strengths") or ""), key=f"edit_strengths_{evaluation_id}")
            concerns = st.text_area("Concerns", value=str(detail.get("concerns") or ""), key=f"edit_concerns_{evaluation_id}")

            selected_mode = st.selectbox(
                "Product Mode",
                list(MODE_LABELS.keys()),
                index=list(MODE_LABELS.keys()).index(mode),
                format_func=lambda x: MODE_LABELS[x],
                key=f"edit_mode_{evaluation_id}",
            )
            submit_edit = st.form_submit_button("Save evaluation changes")

        if submit_edit:
            payload = build_payload_from_form(
                "",
                player_id,
                player_name,
                position,
                age,
                float(offense_rating),
                float(defense_rating),
                float(shooting_rating),
                float(playmaking_rating),
                float(rebounding_rating),
                float(health_risk),
                float(upside),
                float(minutes_stability),
                int(expected_cost_tier),
                team_id,
                timeline,
                float(need_g),
                float(need_f),
                float(need_c),
                float(cap_flexibility),
                float(risk_tolerance),
                summary_note,
                strengths,
                concerns,
                selected_mode,
            )
            try:
                updated = update_evaluation(token, evaluation_id, payload)
                st.session_state["selected_evaluation_id"] = updated.get("id", evaluation_id)
                st.session_state["load_requested"] = True
                st.success(f"Updated evaluation for {player_name}.")
                st.rerun()
            except httpx.HTTPStatusError as exc:
                st.error(f"API error updating evaluation: {exc.response.status_code} {exc.response.text}")
            except Exception as exc:
                st.error(f"Unexpected error updating evaluation: {exc}")


def _csv_cell(row: Dict[str, str], key: str) -> str:
    return (row.get(key) or "").strip()


def _csv_float(row: Dict[str, str], key: str, row_num: int) -> float:
    raw = _csv_cell(row, key)
    try:
        return float(raw)
    except ValueError as exc:
        raise ValueError(f"Row {row_num}: `{key}` must be a number.") from exc


def _csv_int(row: Dict[str, str], key: str, row_num: int) -> int:
    raw = _csv_cell(row, key)
    try:
        return int(float(raw))
    except ValueError as exc:
        raise ValueError(f"Row {row_num}: `{key}` must be a whole number.") from exc


def _payload_to_csv_row(payload: Dict[str, Any]) -> Dict[str, str]:
    player = payload.get("player", {}) or {}
    ctx = payload.get("ctx", {}) or {}
    needs = ctx.get("needs_by_position", {}) or {}
    return {
        "display_name": str(payload.get("display_name") or ""),
        "player_id": str(player.get("id") or ""),
        "player_name": str(player.get("name") or ""),
        "position": str(player.get("position") or ""),
        "age": str(player.get("age") or ""),
        "offense_rating": str(player.get("offense_rating") or ""),
        "defense_rating": str(player.get("defense_rating") or ""),
        "shooting_rating": str(player.get("shooting_rating") or ""),
        "playmaking_rating": str(player.get("playmaking_rating") or ""),
        "rebounding_rating": str(player.get("rebounding_rating") or ""),
        "health_risk": str(player.get("health_risk") or ""),
        "upside": str(player.get("upside") or ""),
        "minutes_stability": str(player.get("minutes_stability") or ""),
        "expected_cost_tier": str(player.get("expected_cost_tier") or ""),
        "team_id": str(ctx.get("team_id") or ""),
        "timeline": str(ctx.get("timeline") or ""),
        "need_g": str(needs.get("G") or ""),
        "need_f": str(needs.get("F") or ""),
        "need_c": str(needs.get("C") or ""),
        "cap_flexibility": str(ctx.get("cap_flexibility") or ""),
        "risk_tolerance": str(ctx.get("risk_tolerance") or ""),
        "summary_note": str(payload.get("summary_note") or ""),
        "strengths": str(payload.get("strengths") or ""),
        "concerns": str(payload.get("concerns") or ""),
        "mode": str(payload.get("mode") or ""),
    }


def build_csv_template_text() -> str:
    buffer = StringIO()
    writer = csv.DictWriter(buffer, fieldnames=CSV_IMPORT_COLUMNS)
    writer.writeheader()
    return buffer.getvalue()


def build_csv_sample_text() -> str:
    sample_rows = [_payload_to_csv_row(payload) for payload in demo_payloads()[:3]]
    buffer = StringIO()
    writer = csv.DictWriter(buffer, fieldnames=CSV_IMPORT_COLUMNS)
    writer.writeheader()
    for row in sample_rows:
        writer.writerow(row)
    return buffer.getvalue()


def parse_csv_import_text(text: str, default_mode: str = DEFAULT_MODE) -> tuple[List[Dict[str, Any]], List[str]]:
    reader = csv.DictReader(StringIO(text))
    payloads: List[Dict[str, Any]] = []
    errors: List[str] = []

    if not reader.fieldnames:
        return [], ["CSV import needs a header row."]

    missing_headers = [col for col in CSV_REQUIRED_COLUMNS if col not in reader.fieldnames]
    if missing_headers:
        return [], [f"Missing required CSV column(s): {', '.join(missing_headers)}"]

    for row_num, row in enumerate(reader, start=2):
        if not any((value or "").strip() for value in row.values()):
            continue
        try:
            for key in CSV_REQUIRED_COLUMNS:
                if not _csv_cell(row, key):
                    raise ValueError(f"Row {row_num}: `{key}` is required.")

            position = _csv_cell(row, "position").upper()
            if position not in {"G", "F", "C"}:
                raise ValueError(f"Row {row_num}: `position` must be G, F, or C.")

            timeline = _csv_cell(row, "timeline")
            if timeline not in {"win_now", "balanced", "rebuild"}:
                raise ValueError(f"Row {row_num}: `timeline` must be win_now, balanced, or rebuild.")

            mode = _csv_cell(row, "mode") or default_mode
            if mode not in MODE_LABELS:
                raise ValueError(f"Row {row_num}: `mode` must be one of {', '.join(MODE_LABELS.keys())}.")

            payloads.append(
                build_payload_from_form(
                    _csv_cell(row, "display_name") or "CSV Import",
                    _csv_cell(row, "player_id"),
                    _csv_cell(row, "player_name"),
                    position,
                    _csv_int(row, "age", row_num),
                    _csv_float(row, "offense_rating", row_num),
                    _csv_float(row, "defense_rating", row_num),
                    _csv_float(row, "shooting_rating", row_num),
                    _csv_float(row, "playmaking_rating", row_num),
                    _csv_float(row, "rebounding_rating", row_num),
                    _csv_float(row, "health_risk", row_num),
                    _csv_float(row, "upside", row_num),
                    _csv_float(row, "minutes_stability", row_num),
                    _csv_int(row, "expected_cost_tier", row_num),
                    _csv_cell(row, "team_id"),
                    timeline,
                    _csv_float(row, "need_g", row_num),
                    _csv_float(row, "need_f", row_num),
                    _csv_float(row, "need_c", row_num),
                    _csv_float(row, "cap_flexibility", row_num),
                    _csv_float(row, "risk_tolerance", row_num),
                    _csv_cell(row, "summary_note"),
                    _csv_cell(row, "strengths"),
                    _csv_cell(row, "concerns"),
                    mode,
                )
            )
        except ValueError as exc:
            errors.append(str(exc))

    return payloads, errors


def split_csv_duplicates(
    payloads: List[Dict[str, Any]], existing_evaluations: List[Dict[str, Any]]
) -> tuple[List[Dict[str, Any]], List[Dict[str, str]], List[Dict[str, Any]]]:
    existing_keys = {
        (
            str(((row.get("player") or {}).get("id") or "")).strip().lower(),
            str(row.get("team_id") or "").strip().lower(),
        ): row
        for row in existing_evaluations
    }

    unique_payloads: List[Dict[str, Any]] = []
    duplicates: List[Dict[str, str]] = []
    duplicate_matches: List[Dict[str, Any]] = []

    for payload in payloads:
        player = payload.get("player", {}) or {}
        ctx = payload.get("ctx", {}) or {}
        key = (
            str(player.get("id") or "").strip().lower(),
            str(ctx.get("team_id") or "").strip().lower(),
        )
        existing = existing_keys.get(key)
        if existing is None:
            unique_payloads.append(payload)
            continue

        duplicates.append(
            {
                "Player": str(player.get("name") or "Player"),
                "Player ID": str(player.get("id") or "—"),
                "Team": str(ctx.get("team_id") or "—"),
                "Existing Score": format_score(existing.get("overall_score")),
                "Existing Mode": MODE_LABELS.get(existing.get("mode") or DEFAULT_MODE, existing.get("mode") or DEFAULT_MODE),
            }
        )
        duplicate_matches.append({"payload": payload, "existing": existing})

    return unique_payloads, duplicates, duplicate_matches


def build_compare_export_markdown(left_detail: Dict[str, Any], right_detail: Dict[str, Any]) -> str:
    left_detail = normalize_detail_for_display(left_detail)
    right_detail = normalize_detail_for_display(right_detail)

    left_player = left_detail.get("player", {}) or {}
    right_player = right_detail.get("player", {}) or {}
    mode = left_detail.get("mode") or right_detail.get("mode") or "pro_wnba"
    roster_need_call = build_roster_need_call(left_detail, right_detail)
    decision_snapshot = build_compare_decision_snapshot(left_detail, right_detail)
    verdicts = build_comparison_verdicts(left_detail, right_detail)

    left_name = left_player.get("name", "Selected Player")
    right_name = right_player.get("name", "Comparison Player")

    lines = [
        f"# {left_name} vs {right_name} - WAIMS-GM Comparison Brief",
        "",
        f"**Mode:** {MODE_LABELS.get(mode, mode)}",
        f"**Primary Call:** {roster_need_call['winner']}",
        "",
        "## Executive Summary",
        compare_summary(left_detail, right_detail),
        "",
        f"## {roster_need_call['title']}",
        roster_need_call["note"],
        "",
        "## Decision Snapshot",
    ]

    if decision_snapshot:
        lines.extend(
            f"- {item['title']}: {item['winner']} - {item['note']}"
            for item in decision_snapshot
        )
    else:
        lines.append("- No decision snapshot available.")

    lines.extend([
        "",
        "## Archetype Fit",
        f"- {build_archetype_comparison_card(left_detail, right_detail)['winner']}: {build_archetype_comparison_card(left_detail, right_detail)['note']}",
        "",
        "## Decision Verdicts",
    ])

    if verdicts:
        lines.extend(
            f"- {item['title']}: {item['winner']} - {item['note']}"
            for item in verdicts
        )
    else:
        lines.append("- No comparison verdicts available.")

    lines.extend(["", "## Component Comparison", f"| Component | {left_name} | Edge | {right_name} |", "|---|---:|---|---:|"])
    for key, label in COMPONENT_LABELS.items():
        left_value = _component_number(left_detail.get("components", {}) or {}, key)
        right_value = _component_number(right_detail.get("components", {}) or {}, key)
        if left_value is None and right_value is None:
            continue

        if left_value is None:
            edge = right_name
        elif right_value is None:
            edge = left_name
        elif abs(left_value - right_value) < 0.01:
            edge = "Even"
        else:
            edge = left_name if left_value > right_value else right_name

        lines.append(f"| {label} | {format_score(left_value)} | {edge} | {format_score(right_value)} |")

    for detail, name in ((left_detail, left_name), (right_detail, right_name)):
        lines.extend(
            [
                "",
                f"## {name} Snapshot",
                f"- Overall Score: {format_score(detail.get('overall_score'))}",
                f"- Recommendation: {clean_action(detail.get('recommended_action'), detail.get('mode'))}",
                f"- Position: {(detail.get('player') or {}).get('position', '—')}",
                f"- Age: {(detail.get('player') or {}).get('age', '—')}",
                f"- Archetype Mix: {summarize_archetype_mix(detail)}",
                f"- Summary Note: {detail.get('summary_note') or 'No summary note.'}",
                "",
                f"### {name} Strengths",
            ]
        )
        strengths = text_block_lines(detail.get("strengths"), "No strengths entered.")
        concerns = text_block_lines(detail.get("concerns"), "No concerns entered.")
        lines.extend(f"- {line}" for line in strengths)
        lines.extend(["", f"### {name} Concerns"])
        lines.extend(f"- {line}" for line in concerns)

    return "\n".join(lines)


def build_export_markdown(detail: Dict[str, Any]) -> str:
    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    mode = detail.get("mode") or "pro_wnba"
    workflow = get_front_office_meta(detail)
    lines = [
        f"# {player.get('name', 'Unknown Player')} — WAIMS-GM Dossier",
        "",
        f"**Mode:** {MODE_LABELS.get(mode, mode)}",
        f"**Recommendation:** {clean_action(detail.get('recommended_action'), mode)}",
        f"**Overall Score:** {format_score(detail.get('overall_score'))}",
        f"**Created:** {format_dt(detail.get('created_at'))}",
        "",
        "## Player Profile",
        f"- Position: {player.get('position', '—')}",
        f"- Age: {player.get('age', '—')}",
        f"- Cost Tier: {player.get('expected_cost_tier', '—')}",
        f"- Health Risk: {player.get('health_risk', '—')}",
        f"- Upside: {player.get('upside', '—')}",
        f"- Minutes Stability: {player.get('minutes_stability', '—')}",
        "",
        "## Executive Memo",
        build_memo(detail),
        "",
        "## Decision Lens",
        build_decision_lens(detail),
        "",
        "## Archetype Mix",
        summarize_archetype_mix(detail),
        "",
        "## Front-Office Workflow",
        f"- Board Stage: {workflow.get('stage', 'Shadow Board')}",
        f"- Priority Tier: {workflow.get('priority_tier', 'Tier 3')}",
        f"- Decision Owner: {workflow.get('owner', 'GM')}",
        f"- Value Posture: {workflow.get('value_posture', 'Price-sensitive')}",
        f"- Market Band: {workflow.get('market_band', 'Market')}",
        f"- Next Action: {workflow.get('next_action', 'No next action entered.')}",
        "",
        f"## {MEDICAL_DILIGENCE_LABEL}",
        f"- Diligence Call: {workflow.get('level', 'Proceed with caution')}",
        f"- Confidence: {workflow.get('confidence', 'Low')}",
        f"- Movement / Functional Flag: {workflow.get('movement_flag', 'No movement review entered.')}",
        f"- Public Injury / Missed-Time History: {workflow.get('public_history', 'No public history summary entered.')}",
        f"- GM-Facing Note: {workflow.get('gm_note', 'No note entered.')}",
        f"- Reviewed By: {workflow.get('reviewed_by', 'No staff review logged yet')}",
        f"- Reviewed At: {format_dt(workflow.get('reviewed_at'))}",
        "",
        "## Research Evidence",
    ]
    evidence_lines = _evidence_lines(str(workflow.get("research_evidence_log") or workflow.get("research_source_links") or ""))
    if evidence_lines:
        lines.extend(f"- {line}" for line in evidence_lines)
    else:
        lines.append("- No public source log saved yet.")

    lines.extend(["", "## Context Summary", summarize_context(ctx), "", "## Strengths"])
    strengths = [x.strip("•- ").strip() for x in (detail.get("strengths") or "").splitlines() if x.strip()]
    concerns = [x.strip("•- ").strip() for x in (detail.get("concerns") or "").splitlines() if x.strip()]
    if strengths:
        lines.extend([f"- {s}" for s in strengths])
    else:
        lines.append("- No strengths entered.")
    lines.extend(["", "## Concerns"])
    if concerns:
        lines.extend([f"- {c}" for c in concerns])
    else:
        lines.append("- No concerns entered.")

    lines.extend(["", "## Five Layer Diagnostic"])
    for row in compute_five_layer_diagnostic(detail):
        lines.append(f"### {row['layer']} ({row['grade']})")
        lines.append(row["note"])
        lines.append("")

    return "\n".join(lines)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor(31, 31, 31)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(10.5)


def add_bullet_list(doc: Document, text_block: Optional[str], fallback: str) -> None:
    lines = [line.strip("•- ").strip() for line in (text_block or "").splitlines() if line.strip()]
    if not lines:
        lines = [fallback]
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(10.5)


def build_export_docx_bytes(detail: Dict[str, Any]) -> bytes:
    if not WORD_EXPORT_AVAILABLE:
        raise RuntimeError(f"Word export unavailable: {WORD_EXPORT_ERROR or 'python-docx not installed'}")

    player = detail.get("player", {}) or {}
    ctx = detail.get("ctx", {}) or {}
    mode = detail.get("mode") or "pro_wnba"
    assumptions = detail.get("assumptions", {}) or {}
    tension_points = detail.get("tension_points", []) or []
    workflow = get_front_office_meta(detail)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run(f"{player.get('name', 'Unknown Player')} — WAIMS-GM Dossier")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(24, 24, 24)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    rr = kicker.add_run("Executive Basketball Briefing")
    rr.italic = True
    rr.font.size = Pt(9.5)
    rr.font.color.rgb = RGBColor(95, 90, 82)

    meta = doc.add_table(rows=2, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    meta.autofit = True

    meta_entries = [
        ("Mode", MODE_LABELS.get(mode, mode)),
        ("Recommendation", clean_action(detail.get("recommended_action"), mode)),
        ("Overall Score", format_score(detail.get("overall_score"))),
        ("Position", str(player.get("position", "—"))),
        ("Age", str(player.get("age", "—"))),
        ("Created", format_dt(detail.get("created_at"))),
    ]
    idx = 0
    for row in meta.rows:
        for cell in row.cells:
            label, value = meta_entries[idx]
            idx += 1
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_shading(cell, "F6F1E7")
            cell.text = ""
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_after = Pt(1)
            r1 = p1.add_run(label)
            r1.bold = True
            r1.font.size = Pt(8.5)
            r1.font.color.rgb = RGBColor(95, 90, 82)
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            r2 = p2.add_run(value)
            r2.font.size = Pt(10.5)

    add_section_heading(doc, "Executive Memo")
    add_body_paragraph(doc, build_memo(detail))

    add_section_heading(doc, "Decision Lens")
    add_body_paragraph(doc, build_decision_lens(detail))

    add_section_heading(doc, "Archetype Mix")
    add_body_paragraph(doc, summarize_archetype_mix(detail))

    add_section_heading(doc, "Player Profile")
    prof = doc.add_table(rows=3, cols=2)
    prof.alignment = WD_TABLE_ALIGNMENT.CENTER
    prof.style = "Table Grid"
    prof_entries = [
        ("Cost Tier", str(player.get("expected_cost_tier", "—"))),
        ("Health Risk", str(player.get("health_risk", "—"))),
        ("Upside", str(player.get("upside", "—"))),
        ("Minutes Stability", str(player.get("minutes_stability", "—"))),
        ("Offense / Defense", f"{player.get('offense_rating', '—')} / {player.get('defense_rating', '—')}"),
        ("Shooting / Playmaking / Rebounding", f"{player.get('shooting_rating', '—')} / {player.get('playmaking_rating', '—')} / {player.get('rebounding_rating', '—')}"),
    ]
    idx = 0
    for row in prof.rows:
        for cell in row.cells:
            label, value = prof_entries[idx]
            idx += 1
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_after = Pt(1)
            r1 = p1.add_run(label)
            r1.bold = True
            r1.font.size = Pt(8.5)
            r1.font.color.rgb = RGBColor(95, 90, 82)
            p2 = cell.add_paragraph()
            r2 = p2.add_run(value)
            r2.font.size = Pt(10.2)

    add_section_heading(doc, "Context Summary")
    add_body_paragraph(doc, summarize_context(ctx))
    needs = ctx.get("needs_by_position", {}) or {}
    if needs:
        add_body_paragraph(doc, "Positional Needs: " + ", ".join(f"{k}: {v}" for k, v in needs.items()))

    add_section_heading(doc, MEDICAL_DILIGENCE_LABEL)
    add_body_paragraph(doc, f"Diligence Call: {workflow.get('level', 'Proceed with caution')}")
    add_body_paragraph(doc, f"Confidence: {workflow.get('confidence', 'Low')}")
    add_body_paragraph(doc, f"Movement / Functional Flag: {workflow.get('movement_flag', 'No movement review entered.')}")
    add_body_paragraph(doc, f"Public Injury / Missed-Time History: {workflow.get('public_history', 'No public history summary entered.')}")
    add_body_paragraph(doc, f"GM-Facing Note: {workflow.get('gm_note', 'No note entered.')}")
    add_body_paragraph(doc, f"Reviewed By: {workflow.get('reviewed_by', 'No staff review logged yet')}")
    add_body_paragraph(doc, f"Reviewed At: {format_dt(workflow.get('reviewed_at'))}")

    add_section_heading(doc, "Research Evidence")
    evidence_lines = _evidence_lines(str(workflow.get("research_evidence_log") or workflow.get("research_source_links") or ""))
    if evidence_lines:
        for line in evidence_lines:
            add_body_paragraph(doc, f"- {line}")
    else:
        add_body_paragraph(doc, "No public source log saved yet.")

    add_section_heading(doc, "Strengths")
    add_bullet_list(doc, detail.get("strengths"), "No strengths entered.")

    add_section_heading(doc, "Concerns")
    add_bullet_list(doc, detail.get("concerns"), "No concerns entered.")

    add_section_heading(doc, "Assumptions")
    if assumptions:
        for k, v in assumptions.items():
            add_body_paragraph(doc, f"{k.replace('_', ' ').title()}: {v}")
    else:
        add_body_paragraph(doc, "No assumptions recorded.")

    add_section_heading(doc, "Tension Points")
    if tension_points:
        for tp in tension_points:
            add_body_paragraph(doc, f"• {tp}")
    else:
        add_body_paragraph(doc, "No major tension points flagged.")

    add_section_heading(doc, "Five Layer Diagnostic")
    for row in compute_five_layer_diagnostic(detail):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{row['layer']} ({row['grade']})")
        r.bold = True
        r.font.size = Pt(11)
        add_body_paragraph(doc, row["note"])

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("WAIMS-GM | Decision-Support Dossier")
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = RGBColor(120, 120, 120)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def main() -> None:
    inject_css()
    render_header()

    if "load_requested" not in st.session_state:
        st.session_state["load_requested"] = False
    if "workspace_authenticated" not in st.session_state:
        st.session_state["workspace_authenticated"] = False
    if "workspace_username" not in st.session_state:
        st.session_state["workspace_username"] = ""
    if "preferred_mode" not in st.session_state:
        st.session_state["preferred_mode"] = DEFAULT_MODE
    if "preferred_mode_selector" not in st.session_state:
        st.session_state["preferred_mode_selector"] = st.session_state["preferred_mode"]
    if "product_mode" not in st.session_state:
        st.session_state["product_mode"] = st.session_state["preferred_mode"]
    if "csv_import_text" not in st.session_state:
        st.session_state["csv_import_text"] = ""
    if "front_office_meta" not in st.session_state:
        st.session_state["front_office_meta"] = {}
    if "workspace_role" not in st.session_state:
        st.session_state["workspace_role"] = "gm"

    if not is_workspace_authenticated():
        render_workspace_login()
        st.stop()

    preferred_mode = st.session_state["preferred_mode"]

    with st.sidebar:
        st.markdown("## Runtime")
        if IS_LIVE_ENV:
            st.error(f"{WAIMS_ENV_LABEL} environment")
        else:
            st.success(f"{WAIMS_ENV_LABEL} environment")
        if WAIMS_DEMO_MODE:
            st.caption("Backend: local in-app demo data")
            st.info("Local demo mode is on. No bearer token or backend is required.")
        else:
            st.caption(f"Backend: {API_BASE_URL}")

        st.markdown("## Access")
        token = st.text_input(
            "Bearer token",
            type="password",
            placeholder="Paste your Supabase access token here" if not WAIMS_DEMO_MODE else "Not required in local demo mode",
            disabled=WAIMS_DEMO_MODE,
        )
        if not WAIMS_DEMO_MODE:
            st.caption("Paste only the raw token, not the word Bearer.")
        else:
            st.caption("Interview-safe local mode: seeded demo dossiers run fully in memory.")

        st.markdown("## Mode Focus")
        preferred_mode = st.selectbox(
            "Session mode",
            list(MODE_LABELS.keys()),
            index=list(MODE_LABELS.keys()).index(preferred_mode),
            format_func=lambda x: MODE_LABELS[x],
            key="preferred_mode_selector",
        )
        if preferred_mode != st.session_state["preferred_mode"]:
            st.session_state["preferred_mode"] = preferred_mode
            st.session_state["product_mode"] = preferred_mode
        preferred_playbook = get_mode_playbook(preferred_mode)
        st.caption(preferred_playbook["headline"])
        st.caption(preferred_playbook["weight_note"])

        st.markdown("## Workspace Role")
        st.caption(f"Signed in as `{current_workspace_user()}`")
        st.success(WORKSPACE_ROLES.get(current_workspace_role(), current_workspace_role()))
        if current_workspace_role() == "gm":
            st.caption("Full front-office workspace: intake, board control, budget discipline, and compare.")
        else:
            st.caption("Constrained collaborator workspace: prospect research and Med Diligence only.")
        if st.button("Sign out"):
            st.session_state["workspace_authenticated"] = False
            st.session_state["workspace_username"] = ""
            st.session_state["workspace_role"] = "gm"
            st.rerun()

        st.markdown("## Filters")
        if is_collaborator_role():
            st.caption("These controls affect Collaborator Home and the dossier tabs inside Player Dossier.")
        else:
            st.caption("These controls affect Board, Player Dossier, and Compare. They do not change Create Evaluation.")
        mode_scope = st.selectbox(
            "Workspace view" if is_collaborator_role() else "Board view",
            ["Focused", "All"] + list(MODE_LABELS.keys()),
            index=0,
            format_func=lambda x: (
                f"Focused: {MODE_LABELS[preferred_mode]}"
                if x == "Focused"
                else ("All modes" if x == "All" else MODE_LABELS[x])
            ),
        )
        stage_filter = st.selectbox("Workflow Stage" if is_collaborator_role() else "Board Stage", ["All"] + WORKFLOW_STAGES, index=0)
        action_filter = st.selectbox("Recommendation", ["All", "draft", "sign", "pass"], index=0)
        sort_by = st.selectbox("Sort by", ["Created", "Score", "Recommendation", "Mode", "Player Name"], index=0)
        descending = st.checkbox("Descending", value=True)
        hide_placeholder = st.checkbox("Hide placeholder/test junk", value=True)
        st.caption("Filter changes apply immediately. Use reload only when you want fresh backend or demo data.")
        load_data = st.button("Reload data")

        st.markdown("## Selected File")
        delete_selected = False
        if current_workspace_role() == "gm":
            delete_selected = st.button("Delete selected file")
        else:
            st.caption("Collaborator roles cannot delete files.")

        st.markdown("## Export Status")
        if WORD_EXPORT_AVAILABLE:
            st.success("Word export ready")
        else:
            st.warning("Word export unavailable")
            st.caption(WORD_EXPORT_ERROR or "Install python-docx to enable Word export.")
    if load_data:
        st.session_state["load_requested"] = True

    if WAIMS_DEMO_MODE:
        st.session_state["load_requested"] = True
        token = token or "demo-local-token"
    else:
        if not token and not st.session_state["load_requested"]:
            st.info("Paste a sandbox bearer token in the sidebar, click 'Refresh board', and the board will populate from Supabase.")
            st.caption("If you want a fast demo board, run scripts\\seed_demo_data.py first.")
            return

        if not token:
            st.warning("A bearer token is required to load the briefing.")
            return

    try:
        raw_evaluations = get_evaluations(token)
    except httpx.HTTPStatusError as e:
        st.error(f"API error loading evaluations: {e.response.status_code} {e.response.text}")
        return
    except Exception as e:
        st.error(f"Unexpected error loading evaluations: {e}")
        return

    sync_front_office_meta(raw_evaluations)
    mode_filter = preferred_mode if mode_scope == "Focused" else mode_scope

    tab_names = (
        ["Collaborator Home", "Player Dossier", "Staff Reports"]
        if is_collaborator_role()
        else ["Create Evaluation", "Board", "Player Dossier", "Staff Reports", "Compare", "Recruiting"]
    )
    tabs = st.tabs(tab_names)
    tab_map = dict(zip(tab_names, tabs))

    if "Create Evaluation" in tab_map:
        with tab_map["Create Evaluation"]:
            st.markdown('<div class="section-kicker">Prospect Intake</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-title">Create New Evaluation</div>', unsafe_allow_html=True)
            st.markdown(
                """
                <div class="intake-shell"><div class="intake-blurb">
                Build a new player file, evaluate it against current team context, and save it directly into the board.
                </div></div>
                """,
                unsafe_allow_html=True,
            )

            if st.session_state.get("product_mode") != preferred_mode:
                st.session_state["product_mode"] = preferred_mode
            selected_mode = st.selectbox(
                "Product Mode",
                list(MODE_LABELS.keys()),
                index=list(MODE_LABELS.keys()).index(preferred_mode),
                format_func=lambda x: MODE_LABELS[x],
                key="product_mode",
            )
            if selected_mode != st.session_state["preferred_mode"]:
                st.session_state["preferred_mode"] = selected_mode
                st.session_state["preferred_mode_selector"] = selected_mode
            preferred_mode = selected_mode
            mode_presets = PRESETS.get(selected_mode, {})
            render_mode_focus_banner(selected_mode, "create", show_label=False)
            with st.expander("Optional: import evaluations from CSV (after demo)", expanded=False):
                st.caption("This is a secondary workflow. For demos, use the seeded board first. Use CSV import when you want to bring in a batch of player files.")
                csv_download_left, csv_download_right = st.columns(2, gap="small")
                with csv_download_left:
                    st.download_button(
                        "Template CSV",
                        data=build_csv_template_text(),
                        file_name="waims_gm_import_template.csv",
                        mime="text/csv",
                        key="download_csv_template",
                        use_container_width=True,
                    )
                with csv_download_right:
                    st.download_button(
                        "Sample CSV",
                        data=build_csv_sample_text(),
                        file_name="waims_gm_import_sample.csv",
                        mime="text/csv",
                        key="download_csv_sample",
                        use_container_width=True,
                    )
                st.caption("Template CSV gives you the schema only. Sample CSV gives you ready-to-import demo rows.")
                if st.button("Load sample rows", key="load_sample_rows"):
                    st.session_state["csv_import_text"] = build_csv_sample_text()
                uploaded_csv = st.file_uploader("Upload evaluation CSV", type=["csv"], key="evaluation_csv_upload")
                csv_text = st.session_state.get("csv_import_text", "")
                if uploaded_csv is not None:
                    try:
                        csv_text = uploaded_csv.getvalue().decode("utf-8-sig")
                        st.session_state["csv_import_text"] = csv_text
                    except UnicodeDecodeError:
                        st.error("CSV must be UTF-8 encoded.")
                if csv_text:
                    csv_payloads, csv_errors = parse_csv_import_text(csv_text, default_mode=selected_mode)
                    if csv_errors:
                        for error in csv_errors:
                            st.error(error)
                    if csv_payloads:
                        importable_payloads, duplicate_rows, duplicate_matches = split_csv_duplicates(csv_payloads, raw_evaluations)
                        preview_rows = [
                            {
                                "Player": payload["player"]["name"],
                                "Mode": MODE_LABELS.get(payload.get("mode") or DEFAULT_MODE, payload.get("mode") or DEFAULT_MODE),
                                "Team": payload["ctx"]["team_id"],
                                "Position": payload["player"]["position"],
                                "Score Inputs": f"O {payload['player']['offense_rating']:.0f} / D {payload['player']['defense_rating']:.0f}",
                            }
                            for payload in csv_payloads
                        ]
                        st.dataframe(preview_rows, use_container_width=True, hide_index=True)
                        replace_duplicates = False
                        if duplicate_rows:
                            st.warning(f"{len(duplicate_rows)} row(s) already match an existing player_id + team_id on the board.")
                            st.dataframe(duplicate_rows, use_container_width=True, hide_index=True)
                            replace_duplicates = st.checkbox(
                                "Replace matching evaluations instead of skipping them",
                                value=False,
                                key="replace_csv_duplicates",
                            )
                        import_button_label = "Import rows"
                        if duplicate_rows and not replace_duplicates:
                            import_button_label = "Import new rows only"
                        elif duplicate_rows and replace_duplicates:
                            import_button_label = "Replace duplicates and import all rows"

                        actionable_payloads = list(importable_payloads)
                        if replace_duplicates:
                            actionable_payloads = actionable_payloads + [item["payload"] for item in duplicate_matches]

                        if csv_payloads and not actionable_payloads:
                            st.info("All uploaded rows already exist on the current board.")
                        if actionable_payloads and st.button(import_button_label, key="import_csv_rows"):
                            created_count = 0
                            replaced_count = 0
                            import_errors: List[str] = []
                            last_new_id: Optional[str] = None
                            if replace_duplicates:
                                for match in duplicate_matches:
                                    existing_row = match["existing"]
                                    try:
                                        delete_evaluation(token, existing_row["id"])
                                        replaced_count += 1
                                    except httpx.HTTPStatusError as exc:
                                        import_errors.append(
                                            f"Replace step ({existing_row.get('id', 'unknown')}): {exc.response.status_code} {exc.response.text}"
                                        )
                                    except Exception as exc:
                                        import_errors.append(f"Replace step ({existing_row.get('id', 'unknown')}): {exc}")

                            for idx, payload in enumerate(actionable_payloads, start=1):
                                try:
                                    created = create_evaluation(token, payload)
                                    created_count += 1
                                    last_new_id = created.get("evaluation_id") or last_new_id
                                except httpx.HTTPStatusError as exc:
                                    import_errors.append(
                                        f"Row {idx} ({payload['player']['name']}): {exc.response.status_code} {exc.response.text}"
                                    )
                                except Exception as exc:
                                    import_errors.append(f"Row {idx} ({payload['player']['name']}): {exc}")

                            if last_new_id:
                                st.session_state["selected_evaluation_id"] = last_new_id
                            st.session_state["load_requested"] = True
                            if created_count:
                                skipped_count = len(duplicate_rows) if not replace_duplicates else 0
                                summary = f"Imported {created_count} evaluation(s) from CSV."
                                if replaced_count:
                                    summary += f" Replaced {replaced_count} matching evaluation(s)."
                                elif skipped_count:
                                    summary += f" Skipped {skipped_count} duplicate row(s)."
                                st.success(summary)
                            for error in import_errors:
                                st.error(error)
                            if created_count:
                                st.rerun()
            preset_name = st.selectbox("Preset", ["Custom"] + list(mode_presets.keys()), index=0, key="preset_name")
            preset = mode_presets.get(preset_name, {})
            team_presets = get_team_context_presets(selected_mode)
            team_preset_name = st.selectbox(
                "Team Context Preset",
                ["Custom"] + list(team_presets.keys()),
                index=0,
                key="team_preset_name",
            )
            team_preset = team_presets.get(team_preset_name, {})
            st.caption(f"Using {MODE_LABELS.get(selected_mode, selected_mode)} presets and recommendation language for this file.")
            if team_preset:
                st.caption(
                    f"Team context preset loaded: {team_preset_name}. "
                    "Use this to avoid re-entering the same roster context on every file."
                )

            with st.form("main_intake_form"):
                st.markdown("**Player identity**")
                c1, c2 = st.columns(2)
                with c1:
                    display_name = st.text_input("Display name", value="Chris")
                    player_id = st.text_input("Player ID", value="p300")
                    player_name = st.text_input("Player name", value="Prospect Wing")
                with c2:
                    position = st.selectbox("Position", ["G", "F", "C"], index=["G", "F", "C"].index(preset.get("position", "F")))
                    age = st.number_input("Age", min_value=16, max_value=45, value=int(preset.get("age", 23)), step=1)
                    expected_cost_tier = st.number_input("Expected cost tier", min_value=0, max_value=10, value=int(preset.get("expected_cost_tier", 2)), step=1)

                st.markdown("**Basketball profile**")
                c3, c4 = st.columns(2)
                with c3:
                    offense_rating = st.slider("Offense", 0, 100, int(preset.get("offense_rating", 74)))
                    defense_rating = st.slider("Defense", 0, 100, int(preset.get("defense_rating", 79)))
                    shooting_rating = st.slider("Shooting", 0, 100, int(preset.get("shooting_rating", 73)))
                with c4:
                    playmaking_rating = st.slider("Playmaking", 0, 100, int(preset.get("playmaking_rating", 64)))
                    rebounding_rating = st.slider("Rebounding", 0, 100, int(preset.get("rebounding_rating", 71)))
                    minutes_stability = st.slider("Minutes stability", 0.0, 1.0, float(preset.get("minutes_stability", 0.72)), 0.01)

                st.markdown("**Risk and projection**")
                c5, c6 = st.columns(2)
                with c5:
                    health_risk = st.slider("Health risk", 0.0, 1.0, float(preset.get("health_risk", 0.22)), 0.01)
                with c6:
                    upside = st.slider("Upside", 0.0, 1.0, float(preset.get("upside", 0.76)), 0.01)

                st.markdown("**Team context**")
                c7, c8 = st.columns(2)
                with c7:
                    team_id = st.text_input("Team ID", value=str(team_preset.get("team_id") or "team-1"))
                    timeline = st.selectbox(
                        "Timeline",
                        ["win_now", "balanced", "rebuild"],
                        index=["win_now", "balanced", "rebuild"].index(str(team_preset.get("timeline") or "balanced")),
                    )
                    cap_flexibility = st.slider("Cap flexibility", 0.0, 1.0, float(team_preset.get("cap_flexibility", 0.60)), 0.01)
                with c8:
                    risk_tolerance = st.slider("Risk tolerance", 0.0, 1.0, float(team_preset.get("risk_tolerance", 0.40)), 0.01)
                    need_g = st.slider("Need at Guard", 0.0, 1.0, float(team_preset.get("need_g", preset.get("need_g", 0.55))), 0.01)
                    need_f = st.slider("Need at Forward", 0.0, 1.0, float(team_preset.get("need_f", preset.get("need_f", 0.80))), 0.01)
                    need_c = st.slider("Need at Center", 0.0, 1.0, float(team_preset.get("need_c", preset.get("need_c", 0.35))), 0.01)

                st.markdown("**Team preset workflow**")
                c_preset_1, c_preset_2 = st.columns([0.7, 1.3], gap="small")
                with c_preset_1:
                    save_team_preset = st.checkbox("Save this team context as a preset", value=False)
                with c_preset_2:
                    save_team_preset_name = st.text_input(
                        "Preset name",
                        value="",
                        placeholder="Example: 2026 roster reset",
                        disabled=not save_team_preset,
                    )

                st.markdown("**Scouting rationale**")
                summary_note = st.text_area("Summary note", value=preset.get("summary_note", ""))
                strengths = st.text_area("Strengths", value=preset.get("strengths", ""))
                concerns = st.text_area("Concerns", value=preset.get("concerns", ""))

                st.markdown("**Board workflow**")
                c9, c10 = st.columns(2)
                with c9:
                    workflow_stage = st.selectbox("Initial board stage", WORKFLOW_STAGES[:-1], index=0)
                with c10:
                    workflow_owner = st.text_input("Decision owner", value="GM")
                workflow_next_action = st.text_input(
                    "Next action",
                    value="Keep live intel moving and tighten the value line before the portal opens.",
                )

                st.markdown("**Staff workflow handoff**")
                st.caption(
                    "New files start with a neutral Med Diligence default. Sport Science or Medical collaborators add staff-authored public-file review later from the Player Dossier."
                )

                submitted = st.form_submit_button("Create evaluation")

            if submitted:
                payload = build_payload_from_form(
                    display_name, player_id, player_name, position, age,
                    float(offense_rating), float(defense_rating), float(shooting_rating),
                    float(playmaking_rating), float(rebounding_rating),
                    float(health_risk), float(upside), float(minutes_stability),
                    int(expected_cost_tier), team_id, timeline,
                    float(need_g), float(need_f), float(need_c),
                    float(cap_flexibility), float(risk_tolerance),
                    summary_note, strengths, concerns, selected_mode,
                )
                try:
                    created = create_evaluation(token, payload)
                    new_id = created.get("evaluation_id")
                    if new_id:
                        st.session_state["selected_evaluation_id"] = new_id
                        if save_team_preset and save_team_preset_name.strip():
                            save_team_context_preset(
                                selected_mode,
                                save_team_preset_name,
                                {
                                    "team_id": team_id,
                                    "timeline": timeline,
                                    "cap_flexibility": cap_flexibility,
                                    "risk_tolerance": risk_tolerance,
                                    "need_g": need_g,
                                    "need_f": need_f,
                                    "need_c": need_c,
                                },
                            )
                        save_front_office_meta(
                            new_id,
                            stage=workflow_stage,
                            owner=workflow_owner.strip() or "GM",
                            next_action=workflow_next_action.strip(),
                            source="Manual / public-file review",
                        )
                    st.session_state["load_requested"] = True
                    st.success(f"Saved evaluation for {player_name}.")
                    st.rerun()
                except httpx.HTTPStatusError as e:
                    st.error(f"API error creating evaluation: {e.response.status_code} {e.response.text}")
                except Exception as e:
                    st.error(f"Unexpected error creating evaluation: {e}")

    evaluations = prepare_evaluations(raw_evaluations, action_filter, hide_placeholder, mode_filter, stage_filter)
    evaluations = sort_evaluations(evaluations, sort_by, descending)

    if "Collaborator Home" in tab_map:
        with tab_map["Collaborator Home"]:
            render_collaborator_home(evaluations, mode_filter, stage_filter, action_filter)

    if delete_selected:
        selected_id = st.session_state.get("selected_evaluation_id")
        if not selected_id:
            st.error("No evaluation selected.")
        else:
            try:
                delete_evaluation(token, selected_id)
                remaining = [row for row in evaluations if row["id"] != selected_id]
                _front_office_store().pop(selected_id, None)
                st.session_state["selected_evaluation_id"] = remaining[0]["id"] if remaining else None
                st.success("Evaluation deleted.")
                st.rerun()
            except httpx.HTTPStatusError as e:
                st.error(f"API error deleting evaluation: {e.response.status_code} {e.response.text}")
            except Exception as e:
                st.error(f"Unexpected error deleting evaluation: {e}")

    detail = None
    selected_id = st.session_state.get("selected_evaluation_id")
    if evaluations:
        existing_ids = {row["id"] for row in evaluations}
        if selected_id not in existing_ids:
            selected_id = evaluations[0]["id"]
            st.session_state["selected_evaluation_id"] = selected_id
    else:
        selected_id = None

    if "Board" in tab_map:
        with tab_map["Board"]:
            render_mode_focus_banner(preferred_mode, "board")
            if current_workspace_role() == "gm":
                st.caption("Board filters live in the sidebar and apply here, in Player Dossier, Staff Reports, and in Compare.")
            else:
                st.caption("Board filters live in the sidebar and apply here, in Player Dossier, and in Staff Reports.")
            render_summary_cards(evaluations)
            st.markdown(f"<div class='filter-note'>Showing {len(evaluations)} evaluation(s) after filters and sorting.</div>", unsafe_allow_html=True)
            st.caption(f"Active board view: {build_board_filter_summary(preferred_mode, mode_scope, stage_filter, action_filter)}.")
            selected_id = render_decision_board(evaluations)
            if selected_id:
                if current_workspace_role() == "gm":
                    st.caption("Use `Player Dossier` for the player file, `Staff Reports` for research and Med Diligence, or `Compare` for side-by-side analysis.")
                    try:
                        board_detail = get_evaluation_detail(token, selected_id)
                        with st.expander("Selected File Workflow", expanded=False):
                            render_front_office_workspace(board_detail)
                            render_budget_scenario_view(board_detail)
                        with st.expander("Edit Evaluation File", expanded=False):
                            render_edit_evaluation_workspace(token, board_detail)
                    except httpx.HTTPStatusError as e:
                        st.error(f"API error loading board file detail: {e.response.status_code} {e.response.text}")
                    except Exception as e:
                        st.error(f"Unexpected error loading board file detail: {e}")
                else:
                    st.caption("Open `Player Dossier` for file context or `Staff Reports` to contribute prospect research and Med Diligence.")

    if selected_id:
        try:
            detail = get_evaluation_detail(token, selected_id)
        except httpx.HTTPStatusError as e:
            st.error(f"API error loading detail: {e.response.status_code} {e.response.text}")
        except Exception as e:
            st.error(f"Unexpected error loading detail: {e}")

    with tab_map["Player Dossier"]:
        st.markdown('<div class="section-kicker">Dossier Workspace</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-title">Player Dossier</div>', unsafe_allow_html=True)
        st.caption(f"Using the current board view: {build_board_filter_summary(preferred_mode, mode_scope, stage_filter, action_filter)}.")

        if not evaluations:
            st.info("No evaluation dossiers are available yet. Create a new evaluation or seed demo files first.")
        else:
            detail = render_file_selector(
                evaluations,
                selected_id,
                token,
                preferred_mode,
                select_key="dossier_select",
                show_exports=True,
            )
            st.caption("This tab is the player file itself: executive summary first, with deeper evaluation layers available below.")
            render_detail(detail, show_diagnostic=True, show_heading=False)

    if "Staff Reports" in tab_map:
        with tab_map["Staff Reports"]:
            st.markdown('<div class="section-kicker">Staff Reporting Workspace</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-title">Staff Reports</div>', unsafe_allow_html=True)
            st.caption(f"Using the current board view: {build_board_filter_summary(preferred_mode, mode_scope, stage_filter, action_filter)}.")

            if not evaluations:
                st.info("No evaluation files are available yet. Create or seed files first, or broaden the current board filters.")
            else:
                detail = render_file_selector(
                    evaluations,
                    selected_id,
                    token,
                    preferred_mode,
                    select_key="staff_reports_select",
                    show_exports=False,
                )
                summary_cards = build_staff_reports_summary(detail) + [build_verified_source_snapshot(detail)]
                render_soft_card_grid(summary_cards, columns_per_row=4, top_margin="0")
                if is_collaborator_role():
                    st.caption("Use this workspace to create or update staff-authored reports for the selected file. Keep all notes public-file, sport-facing, and advisory.")
                    with st.expander("Prospect Research", expanded=True):
                        render_prospect_research_workspace(detail, editable=True)
                    with st.expander("Research Evidence", expanded=False):
                        render_research_evidence_log(detail)
                    with st.expander(MEDICAL_DILIGENCE_LABEL, expanded=True):
                        render_medical_diligence_workspace(detail, editable=True)
                    with st.expander("Completed Reports", expanded=False):
                        render_completed_reports(detail)
                else:
                    st.caption("This workspace is the GM-facing staff readout: research summary, evidence log, Med Diligence, and completed reports for the selected file.")
                    with st.expander("Prospect Research", expanded=True):
                        render_prospect_research_workspace(detail, editable=False)
                    with st.expander("Research Evidence", expanded=False):
                        render_research_evidence_log(detail)
                    with st.expander(MEDICAL_DILIGENCE_LABEL, expanded=True):
                        render_medical_diligence_workspace(detail, editable=False)
                    with st.expander("Completed Reports", expanded=False):
                        render_completed_reports(detail)

    if "Compare" in tab_map:
        with tab_map["Compare"]:
            st.markdown('<div class="section-kicker">Comparison Workspace</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-title">Side-by-Side Compare</div>', unsafe_allow_html=True)
            st.caption(f"Comparison candidates follow the current board view: {build_board_filter_summary(preferred_mode, mode_scope, stage_filter, action_filter)}.")

            if len(evaluations) < 2:
                st.info("Add or seed at least two evaluations to unlock compare mode.")
            else:
                compare_map = {
                    f"{(e.get('player') or {}).get('name', 'Player')} | {MODE_LABELS.get(e.get('mode') or 'pro_wnba', e.get('mode') or 'pro_wnba')} | {format_score(e.get('overall_score'))}": e["id"]
                    for e in evaluations
                }
                compare_labels = list(compare_map.keys())
                primary_label = next((label for label, value in compare_map.items() if value == selected_id), compare_labels[0])
                primary_choice = st.selectbox("Primary dossier", compare_labels, index=compare_labels.index(primary_label), key="compare_primary")
                primary_id = compare_map[primary_choice]
                st.session_state["selected_evaluation_id"] = primary_id

                comparison_candidates = [(label, value) for label, value in compare_map.items() if value != primary_id]
                comparison_labels = [label for label, _ in comparison_candidates]
                compare_choice = st.selectbox("Compare against", comparison_labels, key="compare_secondary")
                compare_id = dict(comparison_candidates)[compare_choice]

                primary_detail = get_evaluation_detail(token, primary_id)
                compare_detail = get_evaluation_detail(token, compare_id)
                render_mode_focus_banner(primary_detail.get("mode") or preferred_mode, "compare", show_label=False)
                if (primary_detail.get("mode") or "") != (compare_detail.get("mode") or ""):
                    st.info("These dossiers come from different product modes. Use the comparison directionally rather than as a pure like-for-like grade.")

                compare_export_md = build_compare_export_markdown(primary_detail, compare_detail)
                compare_name = ((compare_detail.get("player") or {}).get("name", "comparison")).replace(" ", "_")
                primary_name = ((primary_detail.get("player") or {}).get("name", "player")).replace(" ", "_")
                st.download_button(
                    "Download comparison brief (.md)",
                    data=compare_export_md,
                    file_name=f"{primary_name}_vs_{compare_name}_waims_gm_compare.md",
                    mime="text/markdown",
                    key=f"compare_export_{primary_id}",
                )

                render_compare_block(primary_detail, compare_detail)

    if "Recruiting" in tab_map:
        with tab_map["Recruiting"]:
            recruiting_tab()


if __name__ == "__main__":
    main()



